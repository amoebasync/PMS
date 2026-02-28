#!/usr/bin/env python3
"""
GPSトラッキング＆軌跡ビューア機能 設計書 & 機能一覧更新 生成スクリプト
"""
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = os.path.expanduser("~/Downloads/PMS_設計書")


def set_cell_shading(cell, color_hex):
    shading_elm = docx.oxml.parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1E293B")

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F8FAFC")

    return table


def create_design_doc():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic UI'
    font.size = Pt(10)

    # --- 表紙 ---
    for _ in range(4):
        doc.add_paragraph('')

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('PMS Pro\nGPSトラッキング＆軌跡ビューア機能 設計書')
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('株式会社ティラミス 社内向け基幹管理システム')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph('')

    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('文書バージョン', '1.0'),
        ('作成日', '2026年03月01日'),
        ('対象システム', 'PMS Pro（ポスティング管理システム）'),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        for cell in info_table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_page_break()

    # === 1. 概要 ===
    doc.add_heading('1. 概要', level=1)
    doc.add_paragraph(
        '本設計書は、PMS Pro に新規追加した「GPSトラッキング＆軌跡ビューア機能」の設計内容を記述する。'
        '配布員がモバイルアプリで配布作業を行う際のGPS位置情報をリアルタイムで記録し、'
        '管理者がスケジュール画面から配布軌跡を地図上で確認できる機能である。'
    )

    doc.add_heading('1.1 目的', level=2)
    purposes = [
        '配布員の作業状況をリアルタイムでGPS追跡し、配布品質を可視化する',
        '配布軌跡を地図上で再生可能にし、作業の効率分析を支援する',
        '配布開始・終了時に管理者へ自動通知（画面内通知 + ブラウザOS通知）',
        '歩数・距離・カロリー等のフィットネスデータを記録し、パフォーマンス統計を表示する',
        '配布禁止物件のスキップ記録を位置情報と紐付けて管理する',
        '配布完了時に実績枚数と当日報酬をリアルタイムで算出・表示する',
    ]
    for p in purposes:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('1.2 スコープ', level=2)
    scopes = [
        '対象ユーザー：配布員（pms_distributor_session Cookie） / 社内管理者（pms_session Cookie）',
        'モバイルAPI：配布ライフサイクル管理（START → GPS → 進捗 → スキップ → FINISH）',
        '管理者API：軌跡データ取得・リアルタイム監視・通知管理',
        '管理者UI：GPS軌跡ビューアモーダル（Google Maps）・通知ベル・GPS設定',
        '通知方式：ポーリング（30秒間隔） + ブラウザ Notification API',
    ]
    for s in scopes:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('1.3 配布ライフサイクル', level=2)
    doc.add_paragraph(
        '配布員のモバイルアプリにおける配布作業フローは以下の通り：'
    )
    lifecycle = [
        'ログイン → 当日のスケジュール一覧を表示',
        'スケジュール選択 → 配布物一覧を確認',
        'START ボタン押下 → セッション開始（GPS追跡開始、status → DISTRIBUTING）',
        'GPS座標を10秒間隔でサーバーに送信（歩数・距離・カロリーも同時送信）',
        '500ポスト完了ごとに進捗マイルストーン報告',
        '配布禁止物件に遭遇した場合はスキップ記録',
        'FINISH ボタン押下 → チラシ毎の実績枚数入力 → 配布終了（status → COMPLETED）',
        '当日報酬サマリーを表示',
    ]
    for i, step in enumerate(lifecycle, 1):
        doc.add_paragraph(f'{i}. {step}')

    # === 2. データベース設計 ===
    doc.add_heading('2. データベース設計', level=1)

    doc.add_heading('2.1 Enum 定義', level=2)

    doc.add_paragraph('ScheduleStatus（スケジュール状態）— 既存Enumに DISTRIBUTING を追加')
    add_styled_table(doc,
        ['値', '日本語ラベル', '説明'],
        [
            ['UNSTARTED', '未開始', '初期状態'],
            ['IN_PROGRESS', '進行中', '管理者がスケジュールを作成済み'],
            ['DISTRIBUTING', '配布中', '配布員がSTARTを押した（GPS追跡中）'],
            ['COMPLETED', '完了', '配布終了'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph('IncompleteReason（未完了理由）— 新規')
    add_styled_table(doc,
        ['値', '日本語ラベル', '説明'],
        [
            ['AREA_DONE', 'エリア終了', '配布エリア内の全ポストに投函完了（枚数が残っていても正常終了）'],
            ['GIVE_UP', 'ギブアップ', '配布員の体調不良や時間切れ等による中断'],
            ['OTHER', 'その他', 'その他の理由（incompleteNote に詳細記載）'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph('AdminNotificationType（管理者通知種別）— 新規')
    add_styled_table(doc,
        ['値', '日本語ラベル', '説明'],
        [
            ['DISTRIBUTION_START', '配布開始', '配布員がSTARTを押した際に生成'],
            ['DISTRIBUTION_FINISH', '配布完了', '配布員がFINISHを押した際に生成'],
        ]
    )

    # 2.2 テーブル定義
    doc.add_heading('2.2 DistributionSession テーブル（配布セッション）', level=2)
    doc.add_paragraph('テーブル名: distribution_sessions')
    doc.add_paragraph('1スケジュールにつき1セッション（scheduleId は UNIQUE）。配布員がSTARTを押してからFINISHするまでの作業単位。')
    add_styled_table(doc,
        ['カラム名（DB）', 'Prismaフィールド名', 'データ型', 'NULL許可', '説明'],
        [
            ['id', 'id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['schedule_id', 'scheduleId', 'INT (UNIQUE)', '不可', 'FK → DistributionSchedule'],
            ['distributor_id', 'distributorId', 'INT', '不可', 'FK → FlyerDistributor'],
            ['started_at', 'startedAt', 'DATETIME', '不可', '配布開始日時（デバイス時刻）'],
            ['finished_at', 'finishedAt', 'DATETIME', '可', '配布終了日時（null = 配布中）'],
            ['total_steps', 'totalSteps', 'INT', '不可', '総歩数（default: 0）'],
            ['total_distance', 'totalDistance', 'FLOAT', '不可', '総距離（メートル、default: 0）'],
            ['total_calories', 'totalCalories', 'FLOAT', '不可', '総消費カロリー（default: 0）'],
            ['incomplete_reason', 'incompleteReason', 'ENUM', '可', '未完了理由（null = 全数完了）'],
            ['incomplete_note', 'incompleteNote', 'TEXT', '可', '未完了詳細メモ'],
            ['created_at', 'createdAt', 'DATETIME', '不可', '作成日時'],
            ['updated_at', 'updatedAt', 'DATETIME', '不可', '更新日時'],
        ]
    )
    doc.add_paragraph('インデックス: @@index([distributorId, startedAt])')

    doc.add_heading('2.3 GpsPoint テーブル（GPS座標）', level=2)
    doc.add_paragraph('テーブル名: gps_points')
    doc.add_paragraph('10秒ごとに送信される高頻度データ。1セッションあたり約2,880レコード（8時間稼働時）。')
    add_styled_table(doc,
        ['カラム名（DB）', 'Prismaフィールド名', 'データ型', 'NULL許可', '説明'],
        [
            ['id', 'id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['session_id', 'sessionId', 'INT', '不可', 'FK → DistributionSession (CASCADE)'],
            ['latitude', 'latitude', 'FLOAT', '不可', '緯度'],
            ['longitude', 'longitude', 'FLOAT', '不可', '経度'],
            ['accuracy', 'accuracy', 'FLOAT', '可', 'GPS精度（メートル）'],
            ['timestamp', 'timestamp', 'DATETIME', '不可', 'デバイス時刻'],
            ['steps', 'steps', 'INT', '可', '累積歩数'],
            ['distance', 'distance', 'FLOAT', '可', '累積距離（メートル）'],
            ['calories', 'calories', 'FLOAT', '可', '累積消費カロリー'],
            ['created_at', 'createdAt', 'DATETIME', '不可', '作成日時'],
        ]
    )
    doc.add_paragraph('インデックス: @@index([sessionId, timestamp]) — メイン検索パス')

    doc.add_heading('2.4 ProgressEvent テーブル（進捗報告）', level=2)
    doc.add_paragraph('テーブル名: progress_events')
    doc.add_paragraph('配布員が進捗マイルストーン（デフォルト500枚ごと）に達した際の報告データ。')
    add_styled_table(doc,
        ['カラム名（DB）', 'Prismaフィールド名', 'データ型', 'NULL許可', '説明'],
        [
            ['id', 'id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['session_id', 'sessionId', 'INT', '不可', 'FK → DistributionSession (CASCADE)'],
            ['mailbox_count', 'mailboxCount', 'INT', '不可', '累積ポスト数'],
            ['latitude', 'latitude', 'FLOAT', '可', '報告時の緯度'],
            ['longitude', 'longitude', 'FLOAT', '可', '報告時の経度'],
            ['timestamp', 'timestamp', 'DATETIME', '不可', 'デバイス時刻'],
            ['created_at', 'createdAt', 'DATETIME', '不可', '作成日時'],
        ]
    )
    doc.add_paragraph('インデックス: @@index([sessionId, timestamp])')

    doc.add_heading('2.5 SkipEvent テーブル（スキップ記録）', level=2)
    doc.add_paragraph('テーブル名: skip_events')
    doc.add_paragraph('配布禁止物件に遭遇した際のスキップ記録。prohibitedPropertyId は禁止物件DBとのリレーション（任意）。')
    add_styled_table(doc,
        ['カラム名（DB）', 'Prismaフィールド名', 'データ型', 'NULL許可', '説明'],
        [
            ['id', 'id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['session_id', 'sessionId', 'INT', '不可', 'FK → DistributionSession (CASCADE)'],
            ['latitude', 'latitude', 'FLOAT', '不可', 'スキップ位置の緯度'],
            ['longitude', 'longitude', 'FLOAT', '不可', 'スキップ位置の経度'],
            ['prohibited_property_id', 'prohibitedPropertyId', 'INT', '可', 'FK → ProhibitedProperty'],
            ['reason', 'reason', 'TEXT', '可', 'スキップ理由'],
            ['timestamp', 'timestamp', 'DATETIME', '不可', 'デバイス時刻'],
            ['created_at', 'createdAt', 'DATETIME', '不可', '作成日時'],
        ]
    )
    doc.add_paragraph('インデックス: @@index([sessionId])')

    doc.add_heading('2.6 AdminNotification テーブル（管理者通知）', level=2)
    doc.add_paragraph('テーブル名: admin_notifications')
    doc.add_paragraph('配布開始・終了時に自動生成される管理者向け通知。30秒ポーリングで取得。')
    add_styled_table(doc,
        ['カラム名（DB）', 'Prismaフィールド名', 'データ型', 'NULL許可', '説明'],
        [
            ['id', 'id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['type', 'type', 'ENUM', '不可', '通知種別（DISTRIBUTION_START / DISTRIBUTION_FINISH）'],
            ['title', 'title', 'VARCHAR(200)', '不可', '通知タイトル（例: "田中太郎さんが配布を開始しました"）'],
            ['message', 'message', 'TEXT', '可', '通知詳細（例: "西新宿1丁目 / チラシ3種"）'],
            ['schedule_id', 'scheduleId', 'INT', '可', 'FK → DistributionSchedule'],
            ['distributor_id', 'distributorId', 'INT', '可', 'FK → FlyerDistributor'],
            ['is_read', 'isRead', 'BOOLEAN', '不可', '既読フラグ（default: false）'],
            ['created_at', 'createdAt', 'DATETIME', '不可', '作成日時'],
        ]
    )
    doc.add_paragraph('インデックス: @@index([isRead, createdAt]) — 未読通知の取得を高速化')

    # === 3. API設計 ===
    doc.add_heading('3. API設計', level=1)

    doc.add_heading('3.1 スタッフAPI（モバイルアプリ向け）', level=2)
    doc.add_paragraph('認証: pms_distributor_session Cookie（getDistributorFromCookie()）')

    add_styled_table(doc,
        ['メソッド', 'エンドポイント', '説明', '高頻度'],
        [
            ['GET', '/api/staff/config', 'アプリ設定取得（GPS間隔・マイルストーン）', ''],
            ['POST', '/api/staff/distribution/start', '配布セッション開始（status→DISTRIBUTING + 通知）', ''],
            ['POST', '/api/staff/distribution/gps', 'GPS座標受信（10秒ごと）', '⚡'],
            ['POST', '/api/staff/distribution/progress', '進捗マイルストーン報告', ''],
            ['POST', '/api/staff/distribution/skip', '禁止物件スキップ記録', ''],
            ['POST', '/api/staff/distribution/finish', '配布終了（actualCount更新 + status→COMPLETED + 通知 + 報酬計算）', ''],
            ['GET', '/api/staff/distribution/earnings', '当日報酬表示', ''],
        ]
    )

    doc.add_heading('3.1.1 GET /api/staff/config', level=3)
    doc.add_paragraph('SystemSettingテーブルから設定値を取得し、モバイルアプリの動作パラメータを返す。')
    doc.add_paragraph('レスポンス:')
    doc.add_paragraph('{ "gpsTrackingIntervalSeconds": 10, "progressMilestone": 500 }')

    doc.add_heading('3.1.2 POST /api/staff/distribution/start', level=3)
    doc.add_paragraph('配布セッションを開始する。トランザクション内で以下を実行:')
    steps = [
        'DistributionSession を作成（startedAt = デバイス時刻）',
        'DistributionSchedule.status を DISTRIBUTING に更新',
        '初回 GpsPoint を挿入',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('トランザクション後に AdminNotification（DISTRIBUTION_START）を作成。')
    doc.add_paragraph('リクエスト: { "scheduleId": 123, "latitude": 35.658, "longitude": 139.701, "timestamp": "..." }')
    doc.add_paragraph('レスポンス: { "sessionId": 456 }')

    doc.add_heading('3.1.3 POST /api/staff/distribution/gps（⚡高頻度）', level=3)
    doc.add_paragraph(
        '10秒ごとに送信されるGPS座標を受信する。最軽量設計: '
        'セッション所有権チェック（select: { id: true }）+ GpsPoint.create() のみ。'
        'トランザクション不要・監査ログ不要。'
    )
    doc.add_paragraph(
        'フィットネスデータ（steps/distance/calories）が含まれる場合はセッションの総計も更新。'
    )
    doc.add_paragraph('リクエスト: { "sessionId": 456, "latitude": 35.659, "longitude": 139.702, "accuracy": 5.2, "timestamp": "...", "steps": 150, "distance": 120.5, "calories": 8.3 }')
    doc.add_paragraph('レスポンス: { "ok": true }')

    doc.add_heading('3.1.4 POST /api/staff/distribution/progress', level=3)
    doc.add_paragraph('500ポスト（設定可変）ごとの進捗マイルストーン報告。ProgressEvent を作成。')
    doc.add_paragraph('リクエスト: { "sessionId": 456, "mailboxCount": 500, "latitude": ..., "longitude": ..., "timestamp": "..." }')

    doc.add_heading('3.1.5 POST /api/staff/distribution/skip', level=3)
    doc.add_paragraph('配布禁止物件に遭遇した際のスキップ記録。SkipEvent を作成。')
    doc.add_paragraph('リクエスト: { "sessionId": 456, "latitude": ..., "longitude": ..., "prohibitedPropertyId": 99, "reason": "表示あり", "timestamp": "..." }')

    doc.add_heading('3.1.6 POST /api/staff/distribution/finish', level=3)
    doc.add_paragraph('配布を終了する。トランザクション内で以下を実行:')
    steps = [
        '各チラシの DistributionItem.actualCount を更新',
        'DistributionSession を終了（finishedAt、フィットネス総計、未完了理由）',
        'DistributionSchedule.status を COMPLETED に更新',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('トランザクション後に AdminNotification（DISTRIBUTION_FINISH）を作成。')
    doc.add_paragraph('レスポンスに当日報酬サマリー（earnings）を含める。')

    doc.add_heading('3.1.7 GET /api/staff/distribution/earnings', level=3)
    doc.add_paragraph('指定日の報酬を計算して返す。給与計算と同じロジック:')
    doc.add_paragraph('unitPrice = baseRate(rate1Type〜rate6Type) + areaUnitPrice + sizeUnitPrice')
    doc.add_paragraph('earnedAmount = floor(unitPrice × max(actualCounts))')
    doc.add_paragraph('クエリ: ?date=YYYY-MM-DD')

    doc.add_heading('3.2 管理者API', level=2)
    doc.add_paragraph('認証: pms_session Cookie')

    add_styled_table(doc,
        ['メソッド', 'エンドポイント', '説明'],
        [
            ['GET', '/api/schedules/[id]/trajectory', '軌跡データ全取得（GPS・進捗・スキップ・エリア・禁止物件）'],
            ['GET', '/api/schedules/[id]/trajectory/latest', 'リアルタイム最新座標（ポーリング用）'],
            ['GET', '/api/admin/notifications', '通知一覧取得（unreadOnly, limit, sinceId）'],
            ['PUT', '/api/admin/notifications/read', '通知既読（ids指定 or all:true）'],
        ]
    )

    doc.add_heading('3.2.1 GET /api/schedules/[id]/trajectory', level=3)
    doc.add_paragraph('軌跡ビューアに必要な全データを一括取得:')
    items = [
        'session: セッション情報（開始・終了日時、歩数・距離・カロリー、未完了理由）',
        'gpsPoints: 全GPS座標（時系列順）',
        'progressEvents: 進捗マイルストーン一覧',
        'skipEvents: スキップ記録一覧',
        'area: エリア情報（boundary_geojson で地図にポリゴン描画）',
        'prohibitedProperties: エリア内の配布禁止物件一覧',
        'schedule: スケジュール詳細（配布員名、チラシ一覧）',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2.2 GET /api/schedules/[id]/trajectory/latest', level=3)
    doc.add_paragraph('アクティブセッションのリアルタイム監視用。最新GPS座標と最新進捗を返す。')
    doc.add_paragraph('レスポンス: { "latestPoint": { lat, lng, timestamp }, "latestProgress": { mailboxCount }, "isActive": true, "pointCount": 150 }')

    doc.add_heading('3.2.3 GET /api/admin/notifications', level=3)
    doc.add_paragraph('管理者通知一覧を取得。30秒ポーリングで呼び出される。')
    doc.add_paragraph('パラメータ: unreadOnly（未読のみ）, limit（取得件数）, sinceId（差分取得）')

    doc.add_heading('3.2.4 PUT /api/admin/notifications/read', level=3)
    doc.add_paragraph('通知を既読にマーク。')
    doc.add_paragraph('リクエスト: { "ids": [1, 2, 3] } または { "all": true }')

    # === 4. フロントエンド設計 ===
    doc.add_heading('4. フロントエンド設計', level=1)

    doc.add_heading('4.1 GPS軌跡ビューア（TrajectoryViewer）', level=2)
    doc.add_paragraph('ファイル: src/components/schedules/TrajectoryViewer.tsx')
    doc.add_paragraph(
        '/schedules ページのGPSボタンから開くフルスクリーンモーダル。'
        'Google Maps上に配布軌跡を表示し、タイムスライダーで時間経過を再生できる。'
    )

    doc.add_paragraph('構成要素:')
    components = [
        'Google Maps（@react-google-maps/api）: エリアポリゴン、GPS軌跡ポリライン、各種マーカー',
        'タイムスライダー: GPS座標の時系列を0-1000のスライダーで操作',
        '再生コントロール: 再生/一時停止ボタン、速度選択（1x/2x/5x/10x）',
        '統計パネル: 歩数・距離・カロリー・作業時間、時間あたりメトリクス',
        'ライブモード: アクティブセッション時は15秒ポーリングでリアルタイム更新',
    ]
    for c in components:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('マーカー種別:')
    add_styled_table(doc,
        ['マーカー', '色', '説明'],
        [
            ['START', '緑', '配布開始地点'],
            ['FINISH', '赤', '配布終了地点'],
            ['進捗マイルストーン', '緑（ラベル付き）', '500, 1000, 1500...ポスト到達地点'],
            ['スキップ', 'オレンジ', '配布禁止物件スキップ地点'],
            ['禁止物件', '赤（小ドット）', 'エリア内の登録済み禁止物件'],
            ['現在位置', '青（パルス）', 'アクティブセッション時の最新位置'],
        ]
    )

    doc.add_heading('4.2 通知ベル（NotificationBell）', level=2)
    doc.add_paragraph('ファイル: src/components/NotificationBell.tsx')
    doc.add_paragraph('管理者レイアウト（LayoutWrapper.tsx）のヘッダーに配置。')

    features = [
        'ベルアイコン + 未読件数バッジ（赤丸）',
        'クリックでドロップダウン展開（最新20件）',
        '30秒ポーリングで自動更新',
        'ブラウザ Notification API でOS通知（初回に権限リクエスト）',
        '通知クリックで該当スケジュールの軌跡ビューアを開く',
        '「全て既読にする」ボタン',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('4.3 スケジュールページ変更', level=2)
    doc.add_paragraph('ファイル: src/app/schedules/page.tsx')
    doc.add_paragraph('GPSボタンの色分けとステータスバッジの追加:')
    add_styled_table(doc,
        ['スケジュール状態', 'GPSボタン', 'ステータスバッジ'],
        [
            ['UNSTARTED', 'グレー（無効）', '（なし）'],
            ['IN_PROGRESS', 'グレー（無効）', '黄: 進行中'],
            ['DISTRIBUTING', '緑（点滅アニメ）', '緑: 配布中（パルス）'],
            ['COMPLETED', '青', '青: 完了'],
        ]
    )

    doc.add_heading('4.4 GPS設定（/settings）', level=2)
    doc.add_paragraph('ファイル: src/app/settings/page.tsx')
    doc.add_paragraph('全般設定タブ内に「GPSトラッキング設定」セクションを追加:')
    add_styled_table(doc,
        ['設定項目', 'SystemSettingキー', 'デフォルト値', '説明'],
        [
            ['GPS送信間隔', 'gpsTrackingInterval', '10（秒）', 'モバイルアプリのGPS送信頻度'],
            ['進捗マイルストーン', 'progressMilestone', '500（枚）', '進捗報告のポスト数単位'],
        ]
    )

    # === 5. データ量見積もり ===
    doc.add_heading('5. データ量見積もり', level=1)
    add_styled_table(doc,
        ['データ', '計算', '日間', '月間'],
        [
            ['gps_points', '10秒間隔 × 8h = 2,880/セッション × 50人', '144,000行', '4,320,000行'],
            ['progress_events', '500枚ごと × 3回/セッション × 50人', '150行', '4,500行'],
            ['skip_events', '平均5回/セッション × 50人', '250行', '7,500行'],
            ['admin_notifications', '50人 × 2（START+FINISH）', '100行', '3,000行'],
        ]
    )
    doc.add_paragraph('')
    doc.add_paragraph('gps_points テーブルが最も急速に成長するため、90日以上のデータはアーカイブまたは削除を検討する。')

    # === 6. ファイル一覧 ===
    doc.add_heading('6. ファイル一覧', level=1)

    doc.add_heading('6.1 新規作成ファイル', level=2)
    add_styled_table(doc,
        ['ファイルパス', '説明'],
        [
            ['src/app/api/staff/config/route.ts', 'アプリ設定取得API'],
            ['src/app/api/staff/distribution/start/route.ts', '配布開始API + 通知作成'],
            ['src/app/api/staff/distribution/gps/route.ts', 'GPS座標受信API（⚡高頻度）'],
            ['src/app/api/staff/distribution/progress/route.ts', '進捗マイルストーン報告API'],
            ['src/app/api/staff/distribution/skip/route.ts', 'スキップ記録API'],
            ['src/app/api/staff/distribution/finish/route.ts', '配布終了API + 通知作成 + 報酬計算'],
            ['src/app/api/staff/distribution/earnings/route.ts', '当日報酬表示API'],
            ['src/app/api/schedules/[id]/trajectory/route.ts', '軌跡データ全取得API'],
            ['src/app/api/schedules/[id]/trajectory/latest/route.ts', 'リアルタイム最新座標API'],
            ['src/app/api/admin/notifications/route.ts', '通知一覧取得API'],
            ['src/app/api/admin/notifications/read/route.ts', '通知既読API'],
            ['src/components/schedules/TrajectoryViewer.tsx', 'GPS軌跡ビューアモーダル'],
            ['src/components/NotificationBell.tsx', '通知ベルコンポーネント'],
        ]
    )

    doc.add_heading('6.2 修正ファイル', level=2)
    add_styled_table(doc,
        ['ファイルパス', '変更内容'],
        [
            ['prisma/schema.prisma', '5モデル + 2 Enum + ScheduleStatus変更 + リレーション追加'],
            ['src/app/schedules/page.tsx', 'GPSボタンにTrajectoryViewerモーダル接続 + DISTRIBUTINGステータス対応'],
            ['src/app/settings/page.tsx', 'GPS設定セクション追加'],
            ['src/components/LayoutWrapper.tsx', 'ヘッダーにNotificationBell追加'],
            ['CLAUDE.md', 'GPSトラッキング機能セクション追加'],
        ]
    )

    # === 7. セキュリティ ===
    doc.add_heading('7. セキュリティ考慮事項', level=1)
    security = [
        'スタッフAPI: pms_distributor_session Cookie による認証必須（getDistributorFromCookie）',
        'セッション所有権チェック: 全API で distributorId の一致を検証（他人のセッションを操作不可）',
        'スケジュール所有権チェック: START 時に schedule.distributorId の一致を検証',
        '管理者API: pms_session Cookie による認証必須',
        'GPS API（⚡高頻度）: 監査ログは性能影響を考慮し省略（セッション所有権チェックで十分）',
        '通知: ポーリング方式のため WebSocket のセキュリティリスクを回避',
    ]
    for s in security:
        doc.add_paragraph(s, style='List Bullet')

    filepath = os.path.join(OUTPUT_DIR, "PMS_GPSトラッキング＆軌跡ビューア_設計書.docx")
    doc.save(filepath)
    print(f"設計書を生成しました: {filepath}")


# ====================================================================
# 機能一覧更新
# ====================================================================

def update_feature_list():
    filepath = os.path.join(OUTPUT_DIR, "PMS_機能一覧.docx")
    if not os.path.exists(filepath):
        print(f"機能一覧ファイルが見つかりません: {filepath}")
        return

    doc = Document(filepath)

    doc.add_heading('GPSトラッキング＆軌跡ビューア機能', level=1)
    doc.add_paragraph('実装日: 2026年03月01日')

    features = [
        ('GPS配布追跡', '配布員のモバイルアプリからGPS座標を10秒間隔で受信・記録'),
        ('配布ライフサイクル管理', 'START（status→DISTRIBUTING）→ GPS追跡 → FINISH（status→COMPLETED）'),
        ('進捗マイルストーン', '500ポストごとの進捗報告を位置情報付きで記録'),
        ('スキップ記録', '配布禁止物件に遭遇した際のスキップを位置情報・理由付きで記録'),
        ('軌跡ビューア', 'Google Maps上で配布軌跡を再生可能なモーダルビューア（タイムスライダー・速度調整）'),
        ('リアルタイム監視', 'アクティブセッションの現在位置を15秒ポーリングで表示'),
        ('管理者通知', '配布開始・完了時に自動通知（画面内ベル + ブラウザOS通知）'),
        ('報酬計算', '配布完了時に当日報酬サマリーをリアルタイム表示'),
        ('パフォーマンス統計', '歩数・距離・カロリー・時間あたりメトリクスの表示'),
        ('GPS設定', '/settings で GPS送信間隔・進捗マイルストーンを管理者が設定可能'),
    ]

    table = doc.add_table(rows=1 + len(features), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['機能名', '説明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1E293B")

    for ri, (name, desc) in enumerate(features):
        table.rows[ri + 1].cells[0].text = name
        table.rows[ri + 1].cells[1].text = desc
        for ci in range(2):
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
        if ri % 2 == 1:
            for ci in range(2):
                set_cell_shading(table.rows[ri + 1].cells[ci], "F8FAFC")

    doc.save(filepath)
    print(f"機能一覧を更新しました: {filepath}")


if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    create_design_doc()
    update_feature_list()
