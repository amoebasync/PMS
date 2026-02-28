#!/usr/bin/env python3
"""
PMS Pro 機能拡張ロードマップ 文書生成スクリプト
"""
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.expanduser("~/Downloads/PMS_設計書")

# ====================================================================
# ユーティリティ
# ====================================================================

def set_cell_shading(cell, color_hex):
    """セルの背景色を設定"""
    shading_elm = docx.oxml.parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """ヘッダー付きテーブルを追加"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1E293B")

    # データ行
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F8FAFC")

    return table


# ====================================================================
# ロードマップ文書生成
# ====================================================================

def create_roadmap_doc():
    doc = Document()

    # デフォルトフォント設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic UI'
    font.size = Pt(10)

    # =========================================================
    # 表紙
    # =========================================================
    for _ in range(4):
        doc.add_paragraph('')

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('PMS Pro\n機能拡張ロードマップ')
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('株式会社ティラミス 社内向け基幹管理システム')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph('')

    # 文書情報テーブル
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('文書バージョン', '1.0'),
        ('作成日', '2026年02月28日'),
        ('対象システム', 'PMS Pro（ポスティング管理システム）'),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        for cell in info_table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_page_break()

    # =========================================================
    # 1. 概要
    # =========================================================
    doc.add_heading('1. 概要', level=1)
    doc.add_paragraph(
        '本ロードマップは、PMS Pro の既存機能（受注管理、配布スケジュール、CRM、ECポータル、'
        '配布員管理、給与計算、QR分析等）を踏まえ、ポスティング・印刷業の基幹システムおよび'
        'ECサイトの観点から今後追加すべき機能を整理したものである。業務インパクト×実装難易度の'
        '2軸で優先度を設定し、Phase 1〜4 の段階的実装を計画する。'
    )

    # 1.1 優先度定義
    doc.add_heading('1.1 優先度定義', level=2)
    add_styled_table(doc,
        ['優先度', '定義', '実装目安'],
        [
            ['A（最優先）', '業務の根幹に直結。クレーム防止・現場オペレーション改善', 'Phase 1-2'],
            ['B（高）', '業務効率化・売上向上に寄与', 'Phase 2-3'],
            ['C（中）', '顧客満足度向上・競合差別化', 'Phase 3-4'],
            ['D（低）', '将来的な拡張・先進機能', 'Phase 4以降'],
        ]
    )

    doc.add_paragraph('')

    # 1.2 実装フェーズ概要
    doc.add_heading('1.2 実装フェーズ概要', level=2)
    add_styled_table(doc,
        ['フェーズ', '期間目安', '主要機能'],
        [
            ['Phase 1', '1-2ヶ月', 'クレーム管理、配布禁止物件DB、在庫アラート'],
            ['Phase 2', '2-3ヶ月', '配布員モバイルアプリ（MVP）、配布報告書自動生成、入金消込'],
            ['Phase 3', '3-4ヶ月', '配布員アプリv2、印刷校正ワークフロー、経営ダッシュボード、配布員評価制度'],
            ['Phase 4', '4-6ヶ月', 'ルート最適化、リピート発注、反響率トラッキング、AI・外部連携'],
        ]
    )

    doc.add_page_break()

    # =========================================================
    # 2. 優先度A — 業務の根幹に直結（早期着手推奨）
    # =========================================================
    doc.add_heading('2. 優先度A — 業務の根幹に直結（早期着手推奨）', level=1)

    # --- 2.1 配布禁止物件データベース ---
    doc.add_heading('2.1 配布禁止物件データベース', level=2)
    doc.add_paragraph('Phase: 1 / 優先度: A')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: クレーム防止のため「チラシ投函禁止」と指定された物件をDB管理し、'
        '配布員に共有する機能。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('主な機能:')
    features_prohibited = [
        '物件登録（住所 or 地図ピン指定、写真、理由、登録日）',
        'エリアマスタ（Areaモデル）との紐付け（町丁目コード）',
        '配布員アプリへの自動プッシュ（担当エリアの禁止物件一覧を配信）',
        'マップ上に禁止物件ピンを重畳表示（Google Maps + 既存GeoJSON境界）',
        '定期見直しフラグ（1年経過で「再確認」アラート）',
    ]
    for f in features_prohibited:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('DB設計案:')
    add_styled_table(doc,
        ['カラム名', 'データ型', 'NULL', '説明'],
        [
            ['id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['address', 'VARCHAR(500)', '不可', '物件住所'],
            ['lat', 'DECIMAL(10,7)', '可', '緯度'],
            ['lng', 'DECIMAL(10,7)', '可', '経度'],
            ['area_id', 'INT (FK)', '可', 'エリアID → areas.id'],
            ['reason', 'TEXT', '可', '禁止理由'],
            ['photo_url', 'VARCHAR(500)', '可', '写真URL（S3）'],
            ['reported_by', 'INT (FK)', '可', '報告者（社員ID）'],
            ['reported_at', 'DATETIME', '不可', '報告日時'],
            ['is_active', 'BOOLEAN', '不可', '有効フラグ'],
            ['review_due_date', 'DATE', '可', '再確認期限'],
            ['created_at', 'DATETIME', '不可', '作成日時'],
            ['updated_at', 'DATETIME', '不可', '更新日時'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph('既存資産との連携:')
    assets_prohibited = [
        'Areaモデル（GeoJSON境界）が既存。エリア地図に重畳する形で実装可能',
        '@react-google-maps/api が ECポータルで導入済み',
        'S3アップロード基盤が既存（写真保存用）',
    ]
    for a in assets_prohibited:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('注意事項:')
    notes_prohibited = [
        '個人情報保護の観点から、物件特定に最小限の情報のみ保持',
        'マンション名+号室ではなく建物単位で管理を推奨',
        'クレーム管理システムとの連携（「投函禁止」クレーム→自動登録）',
    ]
    for n in notes_prohibited:
        doc.add_paragraph(n, style='List Bullet')

    # --- 2.2 配布員用モバイルアプリ ---
    doc.add_heading('2.2 配布員用モバイルアプリ', level=2)
    doc.add_paragraph('Phase: 2（MVP）→ Phase 3（v2） / 優先度: A')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: 配布員がスマートフォンで配布業務を遂行するための専用アプリ。'
        'GPSトラッキング、配布枚数報告、配布エリアの地図表示を主要機能とする。'
    )
    doc.add_paragraph('')
    doc.add_paragraph(
        '技術選定: React Native (Expo) — 既存TypeScript/React資産を最大活用。iOS/Android両対応。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('MVP機能（Phase 2）:')
    mvp_features = [
        'GPSトラッキング（配布中のリアルタイム位置送信、バッテリー消費を考慮した間欠送信 30秒〜1分間隔）',
        '配布枚数報告（スケジュール別にチラシ種類ごとの実績入力、写真添付）',
        '配布エリア地図表示（GeoJSON境界表示、禁止物件ピン、完了エリアのハイライト）',
        'シフト確認・申請',
        '給与明細閲覧',
    ]
    for f in mvp_features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('v2機能（Phase 3）:')
    v2_features = [
        'プッシュ通知（翌日のスケジュール、緊急連絡、禁止物件更新）',
        'ルート表示（最適化エンジン連携）',
        'オフライン対応（電波の悪い地域での配布に対応）',
        '配布完了写真のGPS・タイムスタンプ付き自動記録',
    ]
    for f in v2_features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('追加API案:')
    add_styled_table(doc,
        ['エンドポイント', 'メソッド', '説明'],
        [
            ['/api/staff/gps', 'POST', '位置情報バッチ送信（[{lat, lng, timestamp, accuracy}]）'],
            ['/api/staff/prohibited-properties', 'GET', '担当エリアの禁止物件取得（?areaId=X）'],
            ['/api/staff/push-token', 'POST', 'プッシュ通知トークン登録'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph('既存資産:')
    assets_mobile = [
        '/staff ポータルの全API（シフト、報告、給与、経費）をそのまま利用可能',
        '/api/staff/schedules/report が報告APIとして存在',
        '配布員認証基盤（pms_distributor_session）が既存',
    ]
    for a in assets_mobile:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('管理画面連携:')
    admin_mobile = [
        '/dispatch にリアルタイムGPS表示パネルを追加',
        '配布員の現在位置と進捗率をマップ上に表示',
    ]
    for a in admin_mobile:
        doc.add_paragraph(a, style='List Bullet')

    # --- 2.3 クレーム管理システム ---
    doc.add_heading('2.3 クレーム管理システム', level=2)
    doc.add_paragraph('Phase: 1 / 優先度: A')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: ポスティング業はクレーム対応が事業継続のカギ。'
        'クレームの記録・対応・分析を一元管理する。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('主な機能:')
    features_complaint = [
        'クレーム登録（物件情報、発生日、内容分類：投函禁止/破損/誤配/遅延/その他）',
        '対応記録（対応者、対応内容、対応日、ステータス：未対応/対応中/解決済）',
        'クレーム→禁止物件DBへの自動連携（「投函禁止」クレームの場合）',
        '顧客・配布員・エリア・スケジュールとの紐付け',
        'クレーム傾向分析（エリア別、配布員別、月別推移）',
    ]
    for f in features_complaint:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('DB設計案 — Complaintテーブル:')
    add_styled_table(doc,
        ['カラム名', 'データ型', 'NULL', '説明'],
        [
            ['id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['type', 'ENUM', '不可', '分類（PROHIBITED/DAMAGE/MISDELIVERY/DELAY/OTHER）'],
            ['property_address', 'VARCHAR(500)', '可', '物件住所'],
            ['area_id', 'INT (FK)', '可', 'エリアID'],
            ['customer_id', 'INT (FK)', '可', '顧客ID'],
            ['distributor_id', 'INT (FK)', '可', '配布員ID'],
            ['schedule_id', 'INT (FK)', '可', 'スケジュールID'],
            ['description', 'TEXT', '不可', 'クレーム内容'],
            ['status', 'ENUM', '不可', '未対応/対応中/解決済'],
            ['resolved_at', 'DATETIME', '可', '解決日時'],
            ['resolved_by', 'INT (FK)', '可', '対応者（社員ID）'],
            ['created_at', 'DATETIME', '不可', '登録日時'],
            ['updated_at', 'DATETIME', '不可', '更新日時'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph('DB設計案 — ComplaintResponseテーブル:')
    add_styled_table(doc,
        ['カラム名', 'データ型', 'NULL', '説明'],
        [
            ['id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['complaint_id', 'INT (FK)', '不可', 'クレームID'],
            ['responder_id', 'INT (FK)', '不可', '対応者（社員ID）'],
            ['content', 'TEXT', '不可', '対応内容'],
            ['created_at', 'DATETIME', '不可', '対応日時'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph(
        'ダッシュボード連携: 未対応クレーム件数をアラートバッジとして / に表示。'
    )

    # --- 2.4 配布報告書の自動生成 ---
    doc.add_heading('2.4 配布報告書の自動生成（顧客向け）', level=2)
    doc.add_paragraph('Phase: 2 / 優先度: A')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: 顧客への納品物として「配布完了報告書」PDFを自動生成。競合との差別化要因。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('主な機能:')
    features_report = [
        '受注単位で配布報告書PDFを自動生成',
        '内容：配布エリア地図（GeoJSON着色）、計画枚数vs実績枚数、配布日程、配布員GPS軌跡（オプション）、QRスキャン実績',
        'ECポータルの /portal/orders/[id] からダウンロード可能',
        '管理者が確認・承認してから顧客に公開するワークフロー',
    ]
    for f in features_report:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('既存資産:')
    assets_report = [
        'PDF生成基盤（/api/documents/[id]/pdf）が既存',
        'QRスキャン分析（/portal/analytics）が既存',
        'GeoJSON地図データが既存',
    ]
    for a in assets_report:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_page_break()

    # =========================================================
    # 3. 優先度B — 業務効率化・売上向上に寄与
    # =========================================================
    doc.add_heading('3. 優先度B — 業務効率化・売上向上に寄与', level=1)

    # --- 3.1 印刷校正ワークフロー ---
    doc.add_heading('3.1 印刷校正ワークフロー', level=2)
    doc.add_paragraph('Phase: 3 / 優先度: B')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: 印刷発注→デザイン校正→顧客承認→入稿のフローをシステム内で完結させる。'
        '現在はメール/LINEで手動運用。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('主な機能:')
    features_proof = [
        'デザインファイルのバージョン管理（v1→v2→v3…）',
        '顧客への校正依頼通知（ECポータル or メール）',
        '顧客の承認/修正指示コメント',
        '承認後に自動的に印刷発注ステータスを更新',
        '校正履歴の保存',
    ]
    for f in features_proof:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('DB設計案:')
    add_styled_table(doc,
        ['カラム名', 'データ型', 'NULL', '説明'],
        [
            ['id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['order_id', 'INT (FK)', '不可', '受注ID'],
            ['order_printing_id', 'INT (FK)', '不可', '印刷依頼ID'],
            ['version', 'INT', '不可', 'バージョン番号'],
            ['file_url', 'VARCHAR(500)', '不可', 'デザインファイルURL（S3）'],
            ['status', 'ENUM', '不可', 'PENDING/APPROVED/REVISION_REQUESTED'],
            ['comment', 'TEXT', '可', '顧客コメント'],
            ['reviewed_at', 'DATETIME', '可', '確認日時'],
            ['created_at', 'DATETIME', '不可', 'アップロード日時'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph(
        'ECポータル連携: /portal/orders/[id] に「校正確認」タブを追加。'
    )

    # --- 3.2 入金消込 ---
    doc.add_heading('3.2 入金消込（売掛金管理）', level=2)
    doc.add_paragraph('Phase: 2 / 優先度: B')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: 請求書発行後の入金確認・消込処理。現在は BillingStatement のステータスが '
        'SENT→PAID だが、部分入金や過入金に対応していない。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('主な機能:')
    features_payment = [
        '入金登録（銀行名、入金日、金額、摘要）',
        '請求書との自動マッチング（金額一致で自動消込提案）',
        '部分入金対応（残高管理）',
        '過入金/未入金アラート',
        '入金一覧＆消込状況レポート',
        '会計ソフト連携用CSV出力',
    ]
    for f in features_payment:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('既存資産: BillingStatement, Payment モデルが存在。')

    # --- 3.3 ルート最適化エンジン ---
    doc.add_heading('3.3 ルート最適化エンジン', level=2)
    doc.add_paragraph('Phase: 4 / 優先度: B')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: 配布員の移動距離を最小化し、効率的な配布ルートを提案する。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('主な機能:')
    features_route = [
        '配布エリア群に対する巡回順序の最適化（TSP近似解法）',
        'Google Directions APIで実際の移動距離/時間を算出',
        '配布員アプリにルート表示（ターンバイターンナビ）',
        '結果の保存と実際のGPS軌跡との比較分析',
    ]
    for f in features_route:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph(
        '注意事項: Google Directions APIのコスト管理が必要。日次のAPI呼び出し上限を設定。'
    )

    # --- 3.4 経営ダッシュボード ---
    doc.add_heading('3.4 経営ダッシュボード（BI機能強化）', level=2)
    doc.add_paragraph('Phase: 3 / 優先度: B')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: 現在のダッシュボードはKPIカード4枚のみ。経営判断に必要な分析を強化する。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('追加ウィジェット:')
    widgets = [
        '売上推移（月次/四半期、前年同月比）',
        '顧客別売上ランキング（上位10社、ABC分析）',
        'エリア別収益性ヒートマップ（単価×配布数÷コスト）',
        '配布員パフォーマンス（配布速度、クレーム率、出勤率）',
        'リード→受注 コンバージョンファネル',
        '季節トレンド分析（業種別の繁忙期パターン）',
        '未回収売掛金エイジング（30日/60日/90日超）',
    ]
    for w in widgets:
        doc.add_paragraph(w, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph(
        '既存資産: Recharts が ECポータル分析で導入済み。Order, BillingStatement, Lead のデータが揃っている。'
    )

    doc.add_page_break()

    # =========================================================
    # 4. 優先度C — 顧客満足度向上・差別化
    # =========================================================
    doc.add_heading('4. 優先度C — 顧客満足度向上・差別化', level=1)

    # --- 4.1 リピート発注機能 ---
    doc.add_heading('4.1 リピート発注機能', level=2)
    doc.add_paragraph('Phase: 4 / 優先度: C')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: 定期配布の顧客が毎回同じ設定で発注する手間を省く。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('主な機能:')
    features_repeat = [
        '過去の発注を「テンプレートとして保存」',
        'ワンクリックでカートに追加（エリア・枚数・配布方法を引き継ぎ）',
        '定期発注スケジュール設定（毎月/隔月/四半期）',
        '発注予定の事前通知メール',
    ]
    for f in features_repeat:
        doc.add_paragraph(f, style='List Bullet')

    # --- 4.2 反響率トラッキング強化 ---
    doc.add_heading('4.2 反響率トラッキング強化', level=2)
    doc.add_paragraph('Phase: 4 / 優先度: C')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: QRコードスキャン分析は実装済みだが、ポスティングの最終成果'
        '（問い合わせ・来店・購入）まで追跡する。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('追加機能:')
    features_tracking = [
        '電話番号トラッキング（チラシ専用番号→着信をAPIで計測）',
        'LP/Webフォーム経由のコンバージョン連携（UTMパラメータ+Webhook）',
        'エリア×チラシデザイン×配布方法のABテスト分析',
        '反響率レポートの顧客ポータル公開',
        'ROI自動計算（配布コスト÷反響数）',
    ]
    for f in features_tracking:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph(
        '既存資産: QrCode, QrScanLog が存在。スキャンデータにコンバージョンイベントを紐付ける拡張。'
    )

    # --- 4.3 配布員評価・ランク制度 ---
    doc.add_heading('4.3 配布員評価・ランク制度', level=2)
    doc.add_paragraph('Phase: 3 / 優先度: C')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: 配布品質のインセンティブ設計。FlyerDistributor にKPIフィールドはあるが、'
        '自動評価の仕組みがない。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('主な機能:')
    features_eval = [
        '月次自動評価スコア算出（出勤率×配布速度×クレーム率×GPS遵守率）',
        'ランク自動昇降（BRONZE→SILVER→GOLD→PLATINUM）',
        'ランクに応じた単価テーブル適用',
        '配布員アプリでの自分のランク・スコア閲覧',
        '管理者向け評価一覧・手動調整機能',
    ]
    for f in features_eval:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph(
        '既存資産: FlyerDistributor に ratePlan, KPIフィールドが既存。'
    )

    # --- 4.4 在庫アラート・自動発注提案 ---
    doc.add_heading('4.4 在庫アラート・自動発注提案', level=2)
    doc.add_paragraph('Phase: 1 / 優先度: C（小規模で即効性あり）')
    doc.add_paragraph('')
    doc.add_paragraph(
        '概要: チラシ在庫の枯渇を防止する。現在は Flyer.stockCount + FlyerTransaction で'
        '入出庫管理しているが、アラートがない。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('主な機能:')
    features_stock = [
        'チラシごとの安全在庫数設定',
        '在庫が閾値を下回ったらダッシュボードにアラート + メール通知',
        '消費ペースから在庫切れ予測日を算出',
        '印刷再発注の提案（過去の発注量・リードタイムから推算）',
    ]
    for f in features_stock:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # =========================================================
    # 5. 優先度D — 将来的な拡張
    # =========================================================
    doc.add_heading('5. 優先度D — 将来的な拡張', level=1)

    # --- 5.1 AIチラシデザイン支援 ---
    doc.add_heading('5.1 AIチラシデザイン支援', level=2)
    doc.add_paragraph(
        'チラシのキャッチコピー提案、レイアウト自動生成（Claude API連携）。'
        '顧客がECポータルからデザイン発注時のオプションとして提供。'
    )

    # --- 5.2 天候連動スケジューリング ---
    doc.add_heading('5.2 天候連動スケジューリング', level=2)
    doc.add_paragraph(
        '天気予報API連携で雨天時の配布スケジュール自動延期・再割当。配布員への自動通知。'
    )

    # --- 5.3 多言語配布員マニュアル ---
    doc.add_heading('5.3 多言語配布員マニュアル（動画+テスト）', level=2)
    doc.add_paragraph(
        '外国人配布員向けの研修コンテンツ管理。動画教材 + 理解度テスト + 合格証発行。'
        '現在の /staff/en ポータルの延長線上。'
    )

    # --- 5.4 外部システム連携 ---
    doc.add_heading('5.4 外部システム連携', level=2)
    integrations = [
        '会計ソフト連携（freee / マネーフォワード API）— 請求書・入金データの自動同期',
        'Slack/LINE WORKS連携 — クレーム・承認通知のチャット配信',
        'Google Sheets連携 — 配布実績のリアルタイムスプレッドシート出力（顧客共有用）',
    ]
    for item in integrations:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =========================================================
    # 6. 実装ロードマップ（タイムライン）
    # =========================================================
    doc.add_heading('6. 実装ロードマップ（タイムライン）', level=1)

    # Phase 1
    doc.add_heading('Phase 1（1-2ヶ月）', level=2)
    add_styled_table(doc,
        ['#', '機能', '優先度', '依存関係'],
        [
            ['1', 'クレーム管理システム', 'A', 'なし（最初に着手推奨）'],
            ['2', '配布禁止物件データベース', 'A', 'クレーム管理と連携'],
            ['3', '在庫アラート', 'C', '小規模で即効性あり。並行実装可能'],
        ]
    )

    doc.add_paragraph('')

    # Phase 2
    doc.add_heading('Phase 2（2-3ヶ月）', level=2)
    add_styled_table(doc,
        ['#', '機能', '優先度', '依存関係'],
        [
            ['4', '配布員モバイルアプリ（MVP）', 'A', '禁止物件DBが先行実装済みであること'],
            ['5', '配布報告書自動生成', 'A', 'GPS データ（アプリ）があると品質向上'],
            ['6', '入金消込', 'B', 'なし'],
        ]
    )

    doc.add_paragraph('')

    # Phase 3
    doc.add_heading('Phase 3（3-4ヶ月）', level=2)
    add_styled_table(doc,
        ['#', '機能', '優先度', '依存関係'],
        [
            ['7', '配布員アプリv2', 'A', 'MVP完了後'],
            ['8', '印刷校正ワークフロー', 'B', 'なし'],
            ['9', '経営ダッシュボード強化', 'B', 'クレーム管理データがあると分析可能'],
            ['10', '配布員評価・ランク制度', 'C', 'GPSデータ・クレームデータが必要'],
        ]
    )

    doc.add_paragraph('')

    # Phase 4
    doc.add_heading('Phase 4（4-6ヶ月）', level=2)
    add_styled_table(doc,
        ['#', '機能', '優先度', '依存関係'],
        [
            ['11', 'ルート最適化エンジン', 'B', '配布員アプリ（GPS・地図）必須'],
            ['12', 'リピート発注機能', 'C', 'なし'],
            ['13', '反響率トラッキング強化', 'C', 'QRコード基盤が既存'],
            ['14', 'AI・外部連携', 'D', '各機能の成熟後'],
        ]
    )

    doc.add_page_break()

    # =========================================================
    # 7. 技術的考慮事項
    # =========================================================
    doc.add_heading('7. 技術的考慮事項', level=1)

    # --- 7.1 配布員アプリの技術スタック ---
    doc.add_heading('7.1 配布員アプリの技術スタック', level=2)
    tech_stack = [
        'フレームワーク: React Native (Expo)',
        '地図: react-native-maps + Google Maps SDK',
        'GPS: expo-location（バックグラウンド位置情報）',
        '通知: expo-notifications + Firebase Cloud Messaging',
        'オフライン: AsyncStorage + SQLite（expo-sqlite）',
        '認証: 既存の pms_distributor_session Cookie ベース認証を Bearer Token 方式に拡張',
    ]
    for item in tech_stack:
        doc.add_paragraph(item, style='List Bullet')

    # --- 7.2 GPS データの設計 ---
    doc.add_heading('7.2 GPS データの設計', level=2)
    gps_design = [
        'バッテリー消費を考慮し、30秒〜1分間隔の間欠送信',
        'バッチ送信（5分ごとにまとめて送信）でネットワーク効率化',
        'GPSログテーブル: distributor_id, lat, lng, accuracy, speed, timestamp',
        '1日あたり約500-1000レコード/配布員。月次でのパーティショニングまたはアーカイブを検討',
        '管理画面でのリアルタイム表示は WebSocket または SSE で実装',
    ]
    for item in gps_design:
        doc.add_paragraph(item, style='List Bullet')

    # --- 7.3 パフォーマンス・スケーラビリティ ---
    doc.add_heading('7.3 パフォーマンス・スケーラビリティ', level=2)
    perf_items = [
        '禁止物件DB: エリア単位でのインデックス。配布員アプリへの配信はエリアIDでフィルタし、差分更新',
        'GPSデータ: 大量レコードのため、ホットデータ（当日〜7日）とコールドデータ（8日以降）でテーブル分離を検討',
        '配布報告書PDF: 非同期ジョブキューで生成（サーバー負荷分散）',
    ]
    for item in perf_items:
        doc.add_paragraph(item, style='List Bullet')

    # =========================================================
    # 保存
    # =========================================================
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = os.path.join(OUTPUT_DIR, 'PMS_機能拡張ロードマップ.docx')
    doc.save(path)
    print(f"[OK] ロードマップ文書を保存: {path}")
    return path


# ====================================================================
# 実行
# ====================================================================
if __name__ == '__main__':
    create_roadmap_doc()
    print("\n完了！")
