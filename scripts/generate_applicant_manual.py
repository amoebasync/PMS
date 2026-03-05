"""
応募者管理・採用フロー 手順書 生成スクリプト
python-docx を使用して Word 文書を生成する
"""

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCREENSHOT_DIR = "/tmp/pms_manual"
OUTPUT_PATH = os.path.expanduser("~/Downloads/PMS_設計書/応募者管理_手順書.docx")

def add_horizontal_line(paragraph):
    """段落の下に水平線を追加"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_background(cell, color_hex):
    """テーブルセルの背景色を設定"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_chapter_heading(doc, chapter_num, title):
    """章タイトルを追加"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"第{chapter_num}章　{title}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # Blue
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    add_horizontal_line(p)
    return p

def add_section_heading(doc, title):
    """節タイトルを追加"""
    p = doc.add_paragraph()
    run = p.add_run(f"■ {title}")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x10, 0x79, 0x54)  # Green
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_step_label(doc, step_num, description):
    """ステップラベルを追加"""
    p = doc.add_paragraph()
    run_step = p.add_run(f"STEP {step_num}  ")
    run_step.font.size = Pt(11)
    run_step.font.bold = True
    run_step.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ステップ番号を青背景で表示するためにテーブルを使用
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_background(cell, '1E40AF')
    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(f"  STEP {step_num}　{description}  ")
    run2.font.size = Pt(11)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(2)

    # Remove table borders except background
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    return table

def add_image(doc, filename, caption=None, width_inches=5.5):
    """画像を追加"""
    img_path = os.path.join(SCREENSHOT_DIR, filename)
    if not os.path.exists(img_path):
        p = doc.add_paragraph(f"[画像なし: {filename}]")
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)

    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cp.runs[0].font.italic = True
        cp.paragraph_format.space_after = Pt(8)

def add_note(doc, text):
    """注意書きを追加"""
    p = doc.add_paragraph()
    run = p.add_run(f"💡 {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)

def add_body_text(doc, text):
    """本文テキストを追加"""
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(4)
    return p


def main():
    doc = Document()

    # ページ余白設定
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ============================================================
    # 表紙
    # ============================================================
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("応募者管理・採用フロー")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("操作手順書")
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run("株式会社ティラミス　PMS Pro")
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run("2026年3月作成")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ============================================================
    # 目次
    # ============================================================
    toc_title = doc.add_paragraph("目　次")
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.runs[0].font.size = Pt(16)
    toc_title.runs[0].font.bold = True

    toc_items = [
        ("第1章", "画面構成と基本操作"),
        ("第2章", "応募者の追加方法"),
        ("　2-1", "方法①　手動登録"),
        ("　2-2", "方法②　面接案内メールの送信"),
        ("第3章", "応募者による面接予約フロー（参考）"),
        ("第4章", "面接後の評価・採用ステータス管理"),
        ("第5章", "研修スロットの割り当て"),
        ("第6章", "配布員登録"),
    ]
    for num, text in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0) if num.startswith("　") else Cm(0)
        run_num = p.add_run(f"{num}　")
        run_num.font.bold = True
        run_num.font.size = Pt(11)
        run_title = p.add_run(text)
        run_title.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ============================================================
    # 第1章: 画面構成と基本操作
    # ============================================================
    add_chapter_heading(doc, "1", "画面構成と基本操作")

    add_body_text(doc, "PMS Pro の応募者管理画面（/applicants）では、面接スロットのカレンダービューと応募者一覧を切り替えながら採用管理を行います。")

    add_section_heading(doc, "カレンダービュー")
    add_body_text(doc, "「カレンダー」タブでは、週単位で面接スロットと予約状況を視覚的に確認できます。緑色が空きスロット、青色が予約済みスロットです。")
    add_image(doc, "01_calendar.png", "▲ カレンダービュー：面接スロットの空き状況を週単位で表示", 5.5)

    add_section_heading(doc, "応募者一覧ビュー")
    add_body_text(doc, "「応募者一覧」タブでは、登録済みの応募者を一覧形式で確認・検索できます。フロー状況（面接待ち／研修待ちなど）と採用ステータスでフィルタリングが可能です。")
    add_image(doc, "02_applicants_list.png", "▲ 応募者一覧タブ：フロー・採用ステータスを一覧表示", 5.5)

    doc.add_page_break()

    # ============================================================
    # 第2章: 応募者の追加方法
    # ============================================================
    add_chapter_heading(doc, "2", "応募者の追加方法")
    add_body_text(doc, "応募者を管理システムに追加する方法は2種類あります。")

    # 2-1 手動登録
    add_section_heading(doc, "方法①　手動登録（管理者が直接入力）")
    add_body_text(doc, "管理者が応募者情報を直接システムに入力して登録します。電話や対面での応募者情報をシステムに取り込む場合に使用します。")

    add_step_label(doc, 1, "「手動登録」ボタンをクリック")
    add_body_text(doc, "応募者管理画面の右上にある緑色の「手動登録」ボタンをクリックします。")
    add_image(doc, "02_applicants_list.png", "▲ 画面右上の「手動登録」ボタンをクリック", 5.5)

    add_step_label(doc, 2, "フォームに情報を入力")
    add_body_text(doc, "登録フォームが表示されます。氏名・メールアドレス・電話番号・職種を入力してください。")
    add_image(doc, "03_manual_reg_empty.png", "▲ 手動登録フォーム（入力前）", 5.0)
    add_image(doc, "04_manual_reg_filled.png", "▲ 必要情報を入力した状態", 5.0)
    add_note(doc, "「面接案内メールを自動送信」チェックボックスをオンにすると、登録完了と同時に応募者へ面接予約リンクのメールが自動送信されます。")

    add_step_label(doc, 3, "登録完了")
    add_body_text(doc, "「登録する」ボタンをクリックすると、応募者が登録されます。チェックボックスがオンの場合は面接案内メールも同時に送信されます。")
    add_image(doc, "05_registration_success.png", "▲ 登録成功：緑色のトースト通知が表示される", 5.5)

    doc.add_page_break()

    # 2-2 面接案内メール送信
    add_section_heading(doc, "方法②　既存応募者への面接案内メール送信")
    add_body_text(doc, "すでに登録済みの応募者に対して、面接日程を自分で選べる予約リンクをメールで送信します。")

    add_step_label(doc, 1, "応募者一覧から対象者の「詳細」をクリック")
    add_body_text(doc, "「応募者一覧」タブから対象の応募者の行にある「詳細」ボタンをクリックします。")
    add_image(doc, "02_applicants_list.png", "▲ 対象応募者の行の「詳細」ボタンをクリック", 5.5)

    add_step_label(doc, 2, "詳細モーダルで「面接案内メールを送信」をクリック")
    add_body_text(doc, "詳細モーダルが表示されます。面接日時が「未設定」の場合、「面接案内メールを送信」ボタンが表示されます。このボタンをクリックしてメールを送信します。")
    add_image(doc, "06_detail_modal.png", "▲ 詳細モーダル：「面接案内メールを送信」ボタン", 5.0)

    add_step_label(doc, 3, "送信完了")
    add_body_text(doc, "メール送信が完了すると、画面右上に緑色のトースト通知が表示されます。応募者のメールアドレスに面接予約リンクが届きます。")
    add_image(doc, "07_invitation_sent.png", "▲ 面接案内メール送信完了のトースト通知", 5.5)

    doc.add_page_break()

    # ============================================================
    # 第3章: 応募者による面接予約フロー
    # ============================================================
    add_chapter_heading(doc, "3", "応募者による面接予約フロー（参考）")
    add_body_text(doc, "本章は、応募者側の操作フローを説明します。管理者は実際に操作する必要はありませんが、応募者からの問い合わせ対応の参考にしてください。")

    add_step_label(doc, 1, "応募者が予約リンクを開く")
    add_body_text(doc, "メールに記載された面接予約リンクを開くと、面接日程選択ページが表示されます。")
    add_image(doc, "08_booking_initial.png", "▲ 面接予約ページ（応募者画面）", 5.5)

    add_step_label(doc, 2, "希望の日付を選択")
    add_body_text(doc, "カレンダーから希望の日付をクリックします。空きスロットがある日付は緑色で表示されます。")
    add_image(doc, "09_date_selected.png", "▲ 日付を選択すると下部に時間帯が表示される", 5.5)

    add_step_label(doc, 3, "時間帯を選択して予約")
    add_body_text(doc, "表示された時間帯の中から希望の枠をクリックして「予約する」ボタンを押します。")
    add_image(doc, "10_slot_selected.png", "▲ 時間帯を選択した状態", 5.5)

    add_step_label(doc, 4, "予約完了")
    add_body_text(doc, "予約が完了すると、確認画面とともに Google Meet リンクが表示されます。応募者にも確認メールが自動送信されます。")
    add_image(doc, "11_booking_complete.png", "▲ 予約完了画面：Google Meet リンクが表示される", 5.5)

    doc.add_page_break()

    # ============================================================
    # 第4章: 面接後の評価・採用ステータス管理
    # ============================================================
    add_chapter_heading(doc, "4", "面接後の評価・採用ステータス管理")
    add_body_text(doc, "面接実施後、管理者が評価スコアを入力し採用ステータスを更新します。")

    add_step_label(doc, 1, "カレンダーで面接予約を確認")
    add_body_text(doc, "カレンダービューで予約済みの面接スロットを確認できます。応募者名が青いブロックで表示されます。")
    add_image(doc, "18_calendar_final.png", "▲ カレンダービュー：予約済みスロットに応募者名が表示される", 5.5)

    add_step_label(doc, 2, "応募者一覧から詳細を開く")
    add_body_text(doc, "「応募者一覧」タブに切り替え、評価する応募者の「詳細」ボタンをクリックします。")
    add_image(doc, "19_applicants_list_final.png", "▲ 応募者一覧：面接日時・フロー・採用ステータスを確認", 5.5)

    add_step_label(doc, 3, "基本情報・面接情報を確認")
    add_body_text(doc, "詳細モーダルの上部で、応募者の基本情報・面接日時・Meet リンクを確認できます。「日程変更」「面接キャンセル」ボタンで変更・取り消しも可能です。")
    add_image(doc, "12_eval_modal_top.png", "▲ 詳細モーダル（上部）：基本情報・面接情報", 5.0)

    add_step_label(doc, 4, "評価スコアを入力・採用ステータスを変更")
    add_body_text(doc, "「面接評価」セクションで各項目（日本語能力・英語能力・コミュニケーション・印象）を1〜5で評価します。「ステータス変更」セクションで採用ステータスを「採用」または「不採用」に変更します。「保存」ボタンをクリックして確定します。")
    add_image(doc, "14_eval_scores_status.png", "▲ 評価スコア入力・ステータス変更セクション", 5.0)
    add_note(doc, "採用ステータスを「採用」に変更すると、採用通知メールが応募者に自動送信されます。")

    doc.add_page_break()

    # ============================================================
    # 第5章: 研修スロットの割り当て
    # ============================================================
    add_chapter_heading(doc, "5", "研修スロットの割り当て")
    add_body_text(doc, "採用が決まった応募者に対して、研修日程を割り当てます。管理者が直接指定する方法と、メールで応募者に自己選択させる方法があります。")

    add_step_label(doc, 1, "研修スロット設定カレンダーを確認")
    add_body_text(doc, "詳細モーダルを下にスクロールすると「研修スロット設定」セクションが表示されます。採用ステータスが「採用」の場合のみ表示されます。「今すぐ指定」を選択すると研修カレンダーが表示されます。")
    add_image(doc, "15_training_calendar.png", "▲ 研修スロット設定：カレンダーで日程を選択", 5.0)

    add_step_label(doc, 2, "研修日を選択")
    add_body_text(doc, "カレンダーから研修日をクリックします。研修スロットがある日付は緑色で表示されます。クリックすると選択中の日付が青くハイライトされ、下部に時間帯が表示されます。")
    add_image(doc, "16_training_date_selected.png", "▲ 日付を選択すると時間帯スロットが表示される", 5.0)

    add_step_label(doc, 3, "時間帯を選択して確定")
    add_body_text(doc, "表示された時間帯をクリックして選択します。選択された枠は紫色にハイライトされます。「保存」ボタンをクリックして確定します。")
    add_image(doc, "17_training_slot_selected.png", "▲ 研修スロット選択（紫色にハイライト）", 5.0)
    add_note(doc, "「後でメールで案内」を選択すると、応募者自身が研修日を選べるリンクのメールを送信できます。")

    doc.add_page_break()

    # ============================================================
    # 第6章: 配布員登録
    # ============================================================
    add_chapter_heading(doc, "6", "配布員登録")
    add_body_text(doc, "研修日程が確定した採用済み応募者を、配布員（FlyerDistributor）として正式にシステムに登録します。")

    add_step_label(doc, 1, "「配布員として登録する」ボタンをクリック")
    add_body_text(doc, "研修日程確定後、詳細モーダルの最下部に「配布員として登録する」ボタンが表示されます。このボタンをクリックして登録フォームを開きます。")
    add_image(doc, "13_training_confirmed_distributor.png", "▲ 研修日程確定後：「配布員として登録する」ボタンが表示", 5.0)
    add_note(doc, "「配布員として登録する」ボタンは、採用ステータスが「採用」かつ研修日程が確定している場合のみ表示されます。")

    add_step_label(doc, 2, "生年月日・支店・スタッフIDを入力")
    add_body_text(doc, "配布員登録に必要な情報（生年月日・所属支店・スタッフID）を入力します。生年月日は初期パスワードになります。応募者の基本情報（氏名・連絡先等）は自動的に引き継がれます。")
    add_note(doc, "登録が完了すると、配布員がモバイルアプリにログインできるようになります。初回パスワードは生年月日（YYYYMMDD形式）です。")

    # ============================================================
    # フッター
    # ============================================================
    doc.add_page_break()
    footer_p = doc.add_paragraph("── 以上 ──")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    footer_p.runs[0].font.size = Pt(11)

    # 保存
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"✅ 手順書を生成しました: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
