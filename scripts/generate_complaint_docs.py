#!/usr/bin/env python3
"""
クレーム管理・配布禁止物件 設計書 & 機能一覧更新 生成スクリプト
"""
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = os.path.expanduser("~/Downloads/PMS_設計書")


def set_cell_shading(cell, color_hex):
    shading_elm = docx.oxml.parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1E293B")

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F8FAFC")

    return table


# ====================================================================
# 1. 設計書
# ====================================================================

def create_design_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic UI'
    font.size = Pt(10)

    # --- 表紙 ---
    for _ in range(4):
        doc.add_paragraph('')

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('PMS Pro\nクレーム管理・配布禁止物件DB 設計書')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    doc.add_paragraph('')
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('株式会社ティラミス  —  2026-02-28')
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()

    # --- 1. 概要 ---
    doc.add_heading('1. 概要', level=1)
    doc.add_paragraph(
        'ポスティング事業においてクレーム発生時に物件を配布禁止として登録し、'
        '今後の配布で誤配を防止する仕組み。既存CSVデータ（約2万件）のインポート、'
        'クレームから禁止物件への自動登録フロー、物件写真のS3アップロード、'
        'ポリラインからGeoJSONへの変換（地図表示用）に対応する。'
    )

    doc.add_heading('1.1 主要機能', level=2)
    features = [
        'クレーム管理（登録・ステータス管理・対応履歴・画像）',
        '配布禁止物件DB（登録・検索・地図表示・CSVインポート）',
        'クレーム → 禁止物件の自動登録フロー',
        '禁止理由マスタ・クレーム種別マスタ',
        'ダッシュボードアラート（未対応クレーム件数）',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    # --- 2. データモデル ---
    doc.add_heading('2. データモデル', level=1)

    doc.add_heading('2.1 Enum', level=2)
    add_styled_table(doc, ['Enum名', '値', '説明'], [
        ['ComplaintStatus', 'UNRESOLVED', '未対応'],
        ['', 'IN_PROGRESS', '対応中'],
        ['', 'RESOLVED', '解決済み'],
    ])

    doc.add_heading('2.2 ProhibitedReason（禁止理由マスタ）', level=2)
    add_styled_table(doc, ['カラム', '型', '説明'], [
        ['id', 'Int (PK)', '自動採番'],
        ['name', 'VarChar(100)', '理由名（クレーム、空き家、危険犬等）'],
        ['sortOrder', 'Int', '表示順（default: 100）'],
        ['isActive', 'Boolean', '有効/無効（default: true）'],
    ])

    doc.add_heading('2.3 ComplaintType（クレーム種別マスタ）', level=2)
    add_styled_table(doc, ['カラム', '型', '説明'], [
        ['id', 'Int (PK)', '自動採番'],
        ['name', 'VarChar(100)', '種別名（投函禁止無視、敷地内侵入等）'],
        ['sortOrder', 'Int', '表示順'],
        ['isActive', 'Boolean', '有効/無効'],
    ])

    doc.add_heading('2.4 ProhibitedProperty（配布禁止物件）', level=2)
    add_styled_table(doc, ['カラム', '型', '説明'], [
        ['id', 'Int (PK)', '自動採番'],
        ['prefectureId', 'Int?', '都道府県ID（prefecturesと紐付け）'],
        ['cityId', 'Int?', '市区町村ID（citiesと紐付け）'],
        ['areaId', 'Int?', 'エリアID（areasと紐付け）'],
        ['postalCode', 'VarChar(8)?', '郵便番号'],
        ['address', 'VarChar(500)', '住所全文（必須）'],
        ['buildingName', 'VarChar(200)?', '建物名'],
        ['roomNumber', 'VarChar(50)?', '部屋番号'],
        ['latitude', 'Float?', '緯度'],
        ['longitude', 'Float?', '経度'],
        ['boundaryGeojson', 'LongText?', 'GeoJSONポリゴン（polylineから変換）'],
        ['customerId', 'Int?', '顧客ID（null=全顧客禁止）'],
        ['prohibitedReasonId', 'Int?', '禁止理由ID'],
        ['reasonDetail', 'Text?', '理由詳細（自由記述）'],
        ['originalCode', 'VarChar(50)?', 'CSV PROHIBITED_CD'],
        ['externalCustomerCode', 'VarChar(20)?', 'CSV CLIENT_CD'],
        ['imageUrls', 'Text?', '画像URL（JSON配列）'],
        ['isActive', 'Boolean', '有効/無効'],
        ['deactivatedAt', 'DateTime?', '無効化日時'],
        ['deactivateReason', 'Text?', '無効化理由'],
        ['complaintId', 'Int?', '登録元クレームID'],
        ['importedAt', 'DateTime?', 'CSVインポート日時'],
    ])

    doc.add_heading('2.5 Complaint（クレーム）', level=2)
    add_styled_table(doc, ['カラム', '型', '説明'], [
        ['id', 'Int (PK)', '自動採番'],
        ['occurredAt', 'Date', '発生日'],
        ['receivedAt', 'DateTime', '受付日時（default: now）'],
        ['complaintTypeId', 'Int?', 'クレーム種別ID'],
        ['customerId', 'Int?', 'クレーム元顧客ID'],
        ['distributorId', 'Int?', '原因配布員ID'],
        ['scheduleId', 'Int?', '関連スケジュールID'],
        ['branchId', 'Int?', '関連支店ID'],
        ['address', 'VarChar(500)', '物件住所'],
        ['buildingName', 'VarChar(200)?', '建物名'],
        ['roomNumber', 'VarChar(50)?', '部屋番号'],
        ['latitude', 'Float?', '緯度'],
        ['longitude', 'Float?', '経度'],
        ['title', 'VarChar(200)', 'クレームタイトル'],
        ['description', 'Text', '詳細内容'],
        ['status', 'ComplaintStatus', 'ステータス（default: UNRESOLVED）'],
        ['assigneeId', 'Int?', '対応担当社員ID'],
        ['imageUrls', 'Text?', '画像URL（JSON配列）'],
    ])

    doc.add_heading('2.6 ComplaintResponse（対応履歴）', level=2)
    add_styled_table(doc, ['カラム', '型', '説明'], [
        ['id', 'Int (PK)', '自動採番'],
        ['complaintId', 'Int (FK)', 'クレームID'],
        ['responderId', 'Int?', '対応社員ID'],
        ['content', 'Text', '対応内容'],
        ['createdAt', 'DateTime', '記録日時'],
    ])

    # --- 3. API ---
    doc.add_heading('3. APIエンドポイント', level=1)

    doc.add_heading('3.1 マスタAPI', level=2)
    add_styled_table(doc, ['メソッド', 'パス', '説明'], [
        ['GET', '/api/prohibited-reasons', '禁止理由一覧'],
        ['POST', '/api/prohibited-reasons', '禁止理由作成'],
        ['PUT', '/api/prohibited-reasons?id=X', '禁止理由更新'],
        ['DELETE', '/api/prohibited-reasons?id=X', '禁止理由削除'],
        ['GET', '/api/complaint-types', 'クレーム種別一覧'],
        ['POST', '/api/complaint-types', 'クレーム種別作成'],
        ['PUT', '/api/complaint-types?id=X', 'クレーム種別更新'],
        ['DELETE', '/api/complaint-types?id=X', 'クレーム種別削除'],
    ])

    doc.add_heading('3.2 配布禁止物件API', level=2)
    add_styled_table(doc, ['メソッド', 'パス', '説明'], [
        ['GET', '/api/prohibited-properties', '一覧（フィルタ・ページネーション）'],
        ['POST', '/api/prohibited-properties', '新規登録'],
        ['PUT', '/api/prohibited-properties?id=X', '更新'],
        ['DELETE', '/api/prohibited-properties?id=X', '論理削除（無効化）'],
        ['GET', '/api/prohibited-properties/{id}', '詳細取得'],
        ['PUT', '/api/prohibited-properties/{id}', '詳細更新'],
        ['POST', '/api/prohibited-properties/{id}/images', '画像アップロード（S3）'],
        ['DELETE', '/api/prohibited-properties/{id}/images', '画像削除'],
        ['POST', '/api/prohibited-properties/import', 'CSVインポート'],
        ['GET', '/api/prohibited-properties/map', '地図用データ取得'],
    ])

    doc.add_heading('3.3 クレーム管理API', level=2)
    add_styled_table(doc, ['メソッド', 'パス', '説明'], [
        ['GET', '/api/complaints', '一覧（フィルタ・ページネーション）'],
        ['POST', '/api/complaints', '新規登録'],
        ['GET', '/api/complaints/{id}', '詳細取得'],
        ['PUT', '/api/complaints/{id}', '更新（ステータス変更含む）'],
        ['GET', '/api/complaints/{id}/responses', '対応履歴一覧'],
        ['POST', '/api/complaints/{id}/responses', '対応記録追加'],
        ['POST', '/api/complaints/{id}/register-prohibited', '禁止物件として登録'],
        ['POST', '/api/complaints/{id}/images', '画像アップロード'],
        ['DELETE', '/api/complaints/{id}/images', '画像削除'],
    ])

    # --- 4. フロントエンド ---
    doc.add_heading('4. フロントエンド画面', level=1)

    doc.add_heading('4.1 クレーム管理画面 (/quality/complaints)', level=2)
    doc.add_paragraph('一覧画面: ステータスフィルタ、支店・期間・キーワード検索、ページネーション')
    doc.add_paragraph('詳細モーダル: 2カラムレイアウト、対応履歴タイムライン、ステータス変更、禁止物件登録')
    doc.add_paragraph('新規登録モーダル: 必須/任意フィールド、顧客/配布員/担当者オートコンプリート検索')

    doc.add_heading('4.2 配布禁止物件画面 (/quality/prohibited-properties)', level=2)
    doc.add_paragraph('3タブ構成: 一覧 / 地図 / CSVインポート')
    doc.add_paragraph('一覧タブ: 都道府県→市区町村連動フィルタ、有効/無効フィルタ、詳細モーダル')
    doc.add_paragraph('地図タブ: Google Maps上に赤マーカー+ポリゴン表示、InfoWindow')
    doc.add_paragraph('CSVインポートタブ: xlsxライブラリでパース、プレビュー、バリデーション、一括登録')

    doc.add_heading('4.3 設定画面 (/settings)', level=2)
    doc.add_paragraph('禁止理由タブ: CRUD（名前、表示順、有効/無効）')
    doc.add_paragraph('クレーム種別タブ: CRUD（名前、表示順、有効/無効）')

    doc.add_heading('4.4 ダッシュボード', level=2)
    doc.add_paragraph('未対応クレームアラートカード: 件数表示、クリックでクレーム管理画面に遷移')

    # --- 5. CSVインポート仕様 ---
    doc.add_heading('5. CSVインポート仕様', level=1)

    doc.add_heading('5.1 CSV列マッピング', level=2)
    add_styled_table(doc, ['CSV列名', 'DBカラム', '変換処理'], [
        ['PROHIBITED_CD', 'originalCode', 'そのまま保存'],
        ['CLIENT_CD', 'customerId', 'Customer.customerCodeで検索。空欄=null（全顧客禁止）'],
        ['POSTAL_CD', 'postalCode', 'そのまま保存'],
        ['ADDRESS_CD', 'areaId/prefectureId/cityId', 'Area.address_codeで検索、先頭2桁→都道府県、5桁→市区町村'],
        ['ADDRESS', 'address', 'そのまま保存（必須）'],
        ['BUILDING_NM', 'buildingName', 'そのまま保存'],
        ['ROOM_NO', 'roomNumber', 'そのまま保存'],
        ['LATITUDE', 'latitude', '数値変換'],
        ['LONGITUDE', 'longitude', '数値変換'],
        ['POLYLINE_PATH', 'boundaryGeojson', '@mapbox/polylineでデコード→GeoJSON Polygon変換'],
        ['REMARK', 'reasonDetail', 'そのまま保存'],
    ])

    doc.add_heading('5.2 ポリラインデコード', level=2)
    doc.add_paragraph(
        'CSVのPOLYLINE_PATH列はGoogleエンコードポリライン形式。'
        'サーバーサイドで@mapbox/polylineライブラリを使用してデコードし、'
        'GeoJSON Polygon形式に変換してboundaryGeojsonカラムに保存する。'
        'GeoJSONの座標順序は[lng, lat]（RFC 7946準拠）。'
    )

    # --- 6. 監査ログ ---
    doc.add_heading('6. 監査ログ', level=1)
    doc.add_paragraph(
        '全CUD操作に対して監査ログを記録。writeAuditLog()をトランザクション内で呼び出す。'
        'CSVインポートは一括操作として1件のログを記録。'
    )
    add_styled_table(doc, ['対象モデル', 'アクション', '備考'], [
        ['ProhibitedReason', 'CREATE/UPDATE/DELETE', 'マスタ変更'],
        ['ComplaintType', 'CREATE/UPDATE/DELETE', 'マスタ変更'],
        ['ProhibitedProperty', 'CREATE/UPDATE/DELETE', 'DELETE=論理削除'],
        ['ProhibitedProperty', 'CREATE (import)', '一括インポート（1ログ）'],
        ['Complaint', 'CREATE/UPDATE/STATUS_CHANGE', 'ステータス変更は別アクション'],
        ['ComplaintResponse', 'CREATE', '対応記録追加'],
    ])

    # --- 保存 ---
    path = os.path.join(OUTPUT_DIR, "PMS_クレーム管理・配布禁止物件_設計書.docx")
    doc.save(path)
    print(f"設計書を保存しました: {path}")


# ====================================================================
# 2. 機能一覧に追加
# ====================================================================

def update_feature_list():
    path = os.path.join(OUTPUT_DIR, "PMS_機能一覧.docx")
    if not os.path.exists(path):
        print(f"機能一覧が見つかりません: {path}")
        return

    doc = Document(path)

    # 最後の表を探す
    tables = doc.tables
    if not tables:
        print("テーブルが見つかりません")
        return

    table = tables[-1]  # 最後の表に追加

    new_rows = [
        ['クレーム管理', 'クレームの登録・ステータス管理・対応履歴・画像管理', '/quality/complaints', '2026-02-28'],
        ['配布禁止物件DB', '禁止物件の登録・検索・地図表示・CSVインポート', '/quality/prohibited-properties', '2026-02-28'],
        ['クレーム→禁止物件登録', 'クレームから配布禁止物件を自動登録', '/quality/complaints 内', '2026-02-28'],
        ['禁止理由マスタ', '配布禁止理由のCRUD管理', '/settings 禁止理由タブ', '2026-02-28'],
        ['クレーム種別マスタ', 'クレーム種別のCRUD管理', '/settings クレーム種別タブ', '2026-02-28'],
        ['未対応クレームアラート', 'ダッシュボードに未対応クレーム件数を表示', '/ (ダッシュボード)', '2026-02-28'],
    ]

    for row_data in new_rows:
        row = table.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.save(path)
    print(f"機能一覧を更新しました: {path}")


if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    create_design_doc()
    update_feature_list()
    print("完了!")
