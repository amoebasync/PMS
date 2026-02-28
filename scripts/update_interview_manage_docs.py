#!/usr/bin/env python3
"""
面接変更・キャンセル機能の設計書更新スクリプト
- PMS_応募者管理機能_設計書.docx にセクション追加
- PMS_機能一覧.docx に項目追加
"""
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DOCS_DIR = os.path.expanduser("~/Downloads/PMS_設計書")

# ──────────────────────────────────────────────
# 1. 応募者管理機能 設計書 更新
# ──────────────────────────────────────────────
def update_applicant_doc():
    path = os.path.join(DOCS_DIR, "PMS_応募者管理機能_設計書.docx")
    doc = docx.Document(path)

    # --- 2.4 Applicant テーブルに managementToken フィールドを追記 ---
    # テーブルを見つける（2.4セクション内のテーブル）
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            cells_text = [c.text.strip() for c in row.cells]
            if "building" in cells_text:
                # このテーブルの後にmanagementTokenの行を追加
                new_row = table.add_row()
                new_row.cells[0].text = "managementToken"
                new_row.cells[1].text = "String?"
                new_row.cells[2].text = "@unique @map(\"management_token\") @db.VarChar(64)"
                new_row.cells[3].text = "面接管理トークン（変更・キャンセル用URL用）"
                # フォントサイズを合わせる
                for cell in new_row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
                break

    # --- セクション10を追加（既存の「9. 今後の拡張候補」の前に挿入） ---
    # 末尾に新セクションを追加する方式
    doc.add_page_break()

    h = doc.add_heading("10. 面接変更・キャンセル機能", level=1)

    doc.add_heading("10.1 概要", level=2)
    doc.add_paragraph(
        "応募者が確認メールに記載されたリンクから、面接日時の変更やキャンセルを自分で行える機能。"
        "トークンベースの認証を採用し、ログイン不要の公開ページとして実装する。"
    )

    doc.add_heading("10.2 フロー", level=2)
    flow_items = [
        "応募者が /apply から応募すると、managementToken（64文字ランダムHEX）が自動生成される",
        "確認メールに「面接の変更・キャンセル」リンク（/apply/manage/{token}）が含まれる",
        "応募者がリンクにアクセスすると、現在の面接情報と変更・キャンセルボタンが表示される",
        "日時変更: 空きスロットから新しい日時を選択 → 旧スロット解放 → 新スロット予約 → 確認メール送信",
        "キャンセル: 確認ダイアログ → スロット解放 → hiringStatus を REJECTED に更新 → キャンセルメール送信",
    ]
    for item in flow_items:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("10.3 制約条件", level=2)
    constraints = [
        "面接当日の変更・キャンセルは不可（前日までに手続きが必要）",
        "flowStatus が INTERVIEW_WAITING かつ hiringStatus が IN_PROGRESS の場合のみ変更可能",
        "キャンセルした場合、hiringStatus は REJECTED に自動更新される（再応募は新規応募ページから）",
        "無効なトークンでアクセスした場合はエラーページを表示",
    ]
    for item in constraints:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("10.4 API 設計", level=2)

    # GET API
    doc.add_heading("GET /api/apply/manage/[token]（面接情報取得）", level=3)
    p = doc.add_paragraph()
    p.add_run("認証: ").bold = True
    p.add_run("不要（トークン認証）")
    p = doc.add_paragraph()
    p.add_run("レスポンス: ").bold = True
    p.add_run("応募者名、言語、面接日時、Google Meet URL、職種名、変更可否フラグ")

    # PUT API
    doc.add_heading("PUT /api/apply/manage/[token]（面接時間変更）", level=3)
    p = doc.add_paragraph()
    p.add_run("認証: ").bold = True
    p.add_run("不要（トークン認証）")
    p = doc.add_paragraph()
    p.add_run("リクエスト: ").bold = True
    p.add_run('{ "newSlotId": number }')
    put_steps = [
        "トークンで応募者を検索、ステータスチェック",
        "面接日の前日以前かチェック",
        "新スロットの空き確認",
        "トランザクション: 旧スロット解放 → 新スロット予約 → Google Meet生成 → 監査ログ",
        "変更確認メール送信（非同期）",
    ]
    for item in put_steps:
        doc.add_paragraph(item, style="List Number")

    # DELETE API
    doc.add_heading("DELETE /api/apply/manage/[token]（面接キャンセル）", level=3)
    p = doc.add_paragraph()
    p.add_run("認証: ").bold = True
    p.add_run("不要（トークン認証）")
    del_steps = [
        "トークンで応募者を検索、ステータスチェック",
        "面接日の前日以前かチェック",
        "トランザクション: スロット解放 → hiringStatus を REJECTED に更新 → 監査ログ",
        "キャンセル確認メール送信（非同期）",
    ]
    for item in del_steps:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("10.5 画面設計", level=2)
    doc.add_heading("/apply/manage/[token]（面接管理ページ）", level=3)
    doc.add_paragraph("公開ページ（認証不要）。応募者の language 設定に基づき ja/en を自動切替。")
    screen_items = [
        "Tiramis ロゴヘッダー + 言語切替ボタン",
        "現在の面接情報カード（職種、日時、Google Meet リンク）",
        "「面接時間を変更する」ボタン → 日付タブ + 時間スロットグリッド → 確認モーダル",
        "「面接をキャンセルする」ボタン → 警告確認モーダル",
        "無効トークン → エラーページ / 変更不可ステータス → 警告メッセージ表示",
    ]
    for item in screen_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("10.6 メール設計", level=2)

    # Table for email specs
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    headers = ["メール種別", "トリガー", "件名（ja）", "件名（en）"]
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    rows_data = [
        ["応募確認メール（改修）", "POST /api/apply", "【Tiramis】面接のご予約を承りました", "[Tiramis] Interview Appointment Confirmed"],
        ["時間変更通知メール", "PUT /api/apply/manage/[token]", "【Tiramis】面接日時が変更されました", "[Tiramis] Interview Rescheduled"],
        ["キャンセル確認メール", "DELETE /api/apply/manage/[token]", "【Tiramis】面接キャンセルのお知らせ", "[Tiramis] Interview Cancelled"],
    ]
    for ri, rd in enumerate(rows_data):
        for ci, val in enumerate(rd):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")
    doc.add_paragraph("応募確認メールには「面接の変更・キャンセル」リンクセクションが追加される。")

    doc.add_heading("10.7 監査ログ", level=2)
    audit_items = [
        "面接時間変更: actorType=SYSTEM, action=UPDATE, targetModel=InterviewSlot, description に旧枠ID・新枠IDを記録",
        "面接キャンセル: actorType=SYSTEM, action=STATUS_CHANGE, targetModel=Applicant, hiringStatus の変更を記録",
    ]
    for item in audit_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("10.8 追加ファイル", level=2)
    files = [
        "src/app/api/apply/manage/[token]/route.ts — 面接管理API（GET/PUT/DELETE）",
        "src/app/apply/manage/[token]/page.tsx — 面接管理ページ（公開）",
    ]
    for f in files:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("10.9 修正ファイル", level=2)
    modified = [
        "prisma/schema.prisma — Applicant に managementToken カラム追加",
        "src/app/api/apply/route.ts — トークン生成・メール引数追加",
        "src/lib/mailer.ts — 確認メールに管理リンク追加 + 変更/キャンセルメール新規追加",
        "src/middleware.ts — /apply/manage パスを公開パスに追加",
        "src/components/LayoutWrapper.tsx — /apply/manage パスでサイドバー非表示",
    ]
    for f in modified:
        doc.add_paragraph(f, style="List Bullet")

    doc.save(path)
    print(f"Updated: {path}")


# ──────────────────────────────────────────────
# 2. 機能一覧 更新
# ──────────────────────────────────────────────
def update_feature_list():
    path = os.path.join(DOCS_DIR, "PMS_機能一覧.docx")
    doc = docx.Document(path)

    # 「14. 応募者管理・面接日程調整機能」セクションを見つける
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "応募者管理" in p.text and p.style.name.startswith("Heading 1"):
            target_idx = i
            break

    if target_idx is None:
        print("Warning: Could not find applicant section in feature list")
        return

    # そのセクション内のテーブルに行を追加
    # テーブルを特定する（14番セクションの直後のテーブル）
    for table in doc.tables:
        has_applicant = False
        for row in table.rows:
            for cell in row.cells:
                if "応募" in cell.text or "面接" in cell.text:
                    has_applicant = True
                    break
            if has_applicant:
                break

        if has_applicant:
            # テーブルに行を追加
            new_row = table.add_row()
            cells_data = [
                "面接変更・キャンセル",
                "Interview Change & Cancel",
                "応募者が確認メールのリンクから面接日時の変更やキャンセルを実行可能。トークンベース認証、ja/en対応。前日までの変更制約あり。",
                "/apply/manage/[token]"
            ]
            for ci, val in enumerate(cells_data):
                if ci < len(new_row.cells):
                    new_row.cells[ci].text = val
                    for p in new_row.cells[ci].paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
            break

    doc.save(path)
    print(f"Updated: {path}")


if __name__ == "__main__":
    update_applicant_doc()
    update_feature_list()
    print("Done!")
