"""
PMS ジョブスケジュール設計書 & SLA設計書 生成スクリプト
"""
import docx
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.expanduser("~/Downloads/PMS_設計書")


def set_cell_shading(cell, color_hex):
    """セルの背景色を設定"""
    shading = docx.oxml.OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, bg_color="2B3A4E"):
    """テーブルヘッダー行のスタイル設定"""
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)


def add_table_row(table, cells_data, align_center=False):
    """テーブルに行を追加"""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(text)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
            if align_center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return row


def set_table_style(table):
    """テーブルの基本スタイル"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)


def add_title_page(doc, title, subtitle):
    """表紙ページ"""
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Tiramis Co., Ltd.  |  2026年3月4日")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()


# ============================================================
# 1. ジョブスケジュール設計書
# ============================================================
def generate_job_schedule_doc():
    doc = docx.Document()

    # デフォルトフォント設定
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Yu Gothic"
    font.size = Pt(10)

    add_title_page(doc, "PMS ジョブスケジュール", "設計書")

    # --- 1. 概要 ---
    doc.add_heading("1. 概要", level=1)
    doc.add_paragraph(
        "本書では、PMS（ポスティング管理システム）で定期実行されるバッチジョブ（CRONジョブ）の"
        "一覧・実行スケジュール・認証方式・監視方針・障害時対応を定義する。"
    )
    doc.add_paragraph(
        "すべてのジョブは EC2 の crontab から curl で Next.js API を呼び出す方式で実行される。"
        "認証には Bearer トークン（CRON_SECRET 環境変数）を使用する。"
    )

    # --- 2. アーキテクチャ ---
    doc.add_heading("2. ジョブ実行アーキテクチャ", level=1)
    doc.add_paragraph(
        "■ 実行フロー\n"
        "  crontab (EC2)\n"
        "    → curl -H 'Authorization: Bearer $CRON_SECRET' http://localhost:3000/api/cron/{job}\n"
        "    → Next.js API Route Handler\n"
        "    → Prisma → MySQL (RDS)\n"
        "    → レスポンス → ログファイル (/tmp/pms-cron-*.log)"
    )

    doc.add_heading("2.1 認証方式", level=2)
    doc.add_paragraph(
        "CRON APIエンドポイントは middleware.ts で公開パスとして設定されており、"
        "管理者セッション認証は不要。代わりに各APIハンドラ内で Authorization ヘッダの "
        "Bearer トークンを CRON_SECRET 環境変数と照合する。"
    )

    doc.add_heading("2.2 実行環境", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0]
    hdr.cells[0].text = "項目"
    hdr.cells[1].text = "値"
    style_header_row(hdr)
    add_table_row(table, ["実行ホスト", "EC2 #1 (ip-172-31-33-73), EC2 #2 (ip-172-31-12-171)"])
    add_table_row(table, ["タイムゾーン", "UTC（EC2デフォルト）"])
    add_table_row(table, ["呼出先", "http://localhost:3000/api/cron/*"])
    add_table_row(table, ["ログ出力先", "/tmp/pms-cron-{ジョブ名}.log"])
    add_table_row(table, ["登録スクリプト", "scripts/after_install.sh（CodeDeploy AfterInstall フック）"])
    add_table_row(table, ["重複防止", "crontab -l | grep で既存チェック → 登録済みならスキップ"])
    set_table_style(table)

    # --- 3. ジョブ一覧 ---
    doc.add_heading("3. ジョブ一覧", level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0]
    headers = ["No.", "ジョブ名", "CRON式 (UTC)", "JST", "エンドポイント", "説明"]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    style_header_row(hdr)

    jobs = [
        ("1", "面接スロット自動生成", "0 15 * * *", "00:00", "/api/cron/generate-slots",
         "面接スロットマスタに基づき14日分のスロットを自動生成"),
        ("2", "定期タスク自動生成", "0 15 * * *", "00:00", "/api/cron/generate-tasks",
         "タスクテンプレートに基づきタスクを自動生成（SHARED/INDIVIDUAL展開）"),
        ("3", "研修スロット自動生成", "0 15 * * *", "00:00", "/api/cron/generate-training-slots",
         "研修スロットマスタに基づき14日分を自動生成"),
        ("4", "配布員評価", "0 15 * * 1", "00:00 月曜", "/api/cron/evaluate-distributors",
         "週次の配布員スコア算出・ランク更新"),
        ("5", "ビザ期限チェック", "0 15 * * *", "00:00", "/api/cron/check-visa-expiry",
         "在留期限が近い配布員を検出しアラート生成"),
        ("6", "アラート定義チェック", "0 15 * * *", "00:00", "/api/cron/check-alert-definitions",
         "アラート定義に基づき条件判定・通知生成"),
        ("7", "ハウスキープ", "0 15 * * *", "00:00", "/api/cron/housekeep",
         "通知30日超削除 / 監査ログ90日超S3アーカイブ / GPS 365日超S3アーカイブ"),
    ]

    for job in jobs:
        add_table_row(table, job)
    set_table_style(table)

    # --- 4. 各ジョブ詳細 ---
    doc.add_heading("4. 各ジョブ詳細", level=1)

    # 4.1
    doc.add_heading("4.1 面接スロット自動生成 (generate-slots)", level=2)
    doc.add_paragraph("対象テーブル: interview_slots, interview_slot_masters, default_interview_slots")
    doc.add_paragraph("処理内容:")
    doc.add_paragraph("有効なスロットマスタごとに、曜日設定に従い14日先までのスロットを生成する。"
                      "既に存在するスロット（同日時・同マスタ）は重複生成しない。", style="List Bullet")
    doc.add_paragraph("推定実行時間: 1〜3秒", style="List Bullet")
    doc.add_paragraph("DB負荷: 低（INSERT数十件/回）", style="List Bullet")

    # 4.2
    doc.add_heading("4.2 定期タスク自動生成 (generate-tasks)", level=2)
    doc.add_paragraph("対象テーブル: tasks, task_assignees, task_templates")
    doc.add_paragraph("処理内容:")
    doc.add_paragraph("有効なテンプレートのうち、当日が実行日に該当するものを抽出。"
                      "SHARED: 1タスク＋複数アサイニー。INDIVIDUAL: 部署/支店→社員展開し人数分のタスク生成。"
                      "同日・同テンプレートからの二重生成チェックあり。", style="List Bullet")
    doc.add_paragraph("推定実行時間: 1〜5秒", style="List Bullet")
    doc.add_paragraph("DB負荷: 低〜中（テンプレート数・社員数に依存）", style="List Bullet")

    # 4.3
    doc.add_heading("4.3 研修スロット自動生成 (generate-training-slots)", level=2)
    doc.add_paragraph("対象テーブル: training_slots, default_training_slots")
    doc.add_paragraph("処理内容:")
    doc.add_paragraph("曜日ごとの研修スロットマスタに基づき14日分を生成。"
                      "重複チェック（同日時）あり。", style="List Bullet")
    doc.add_paragraph("推定実行時間: 1秒以下", style="List Bullet")
    doc.add_paragraph("DB負荷: 低", style="List Bullet")

    # 4.4
    doc.add_heading("4.4 配布員評価 (evaluate-distributors)", level=2)
    doc.add_paragraph("対象テーブル: flyer_distributors, distribution_schedules, complaints, distributor_evaluations")
    doc.add_paragraph("処理内容:")
    doc.add_paragraph("直近1週間の出勤日数・配布枚数・クレーム減点を集計し、スコア算出。"
                      "スコアに基づきランク（S/A/B/C/D）を更新。評価履歴を記録。", style="List Bullet")
    doc.add_paragraph("推定実行時間: 5〜30秒（配布員数に依存）", style="List Bullet")
    doc.add_paragraph("DB負荷: 中（全配布員分のSELECT + UPDATE）", style="List Bullet")
    doc.add_paragraph("実行頻度: 週1回（月曜）", style="List Bullet")

    # 4.5
    doc.add_heading("4.5 ビザ期限チェック (check-visa-expiry)", level=2)
    doc.add_paragraph("対象テーブル: flyer_distributors, admin_notifications")
    doc.add_paragraph("処理内容:")
    doc.add_paragraph("在留期限が30日/60日/90日以内の配布員を検出し、管理者通知を生成。", style="List Bullet")
    doc.add_paragraph("推定実行時間: 1〜3秒", style="List Bullet")
    doc.add_paragraph("DB負荷: 低", style="List Bullet")

    # 4.6
    doc.add_heading("4.6 アラート定義チェック (check-alert-definitions)", level=2)
    doc.add_paragraph("対象テーブル: alert_definitions, alert_instances, admin_notifications")
    doc.add_paragraph("処理内容:")
    doc.add_paragraph("アラート定義の条件（SQL/ロジック）を評価し、条件該当時にアラートインスタンスと"
                      "管理者通知を生成。ブラウザ通知フラグに応じてOS通知用データも作成。", style="List Bullet")
    doc.add_paragraph("推定実行時間: 3〜10秒（定義数・条件の複雑さに依存）", style="List Bullet")
    doc.add_paragraph("DB負荷: 中（複数テーブルの条件判定クエリ）", style="List Bullet")

    # 4.7
    doc.add_heading("4.7 ハウスキープ (housekeep)", level=2)
    doc.add_paragraph("対象テーブル: admin_notifications, audit_logs, gps_points")
    doc.add_paragraph("処理内容:")
    doc.add_paragraph("(a) admin_notifications: 30日超の通知を DELETE", style="List Bullet")
    doc.add_paragraph("(b) audit_logs: 90日超のログを S3 に JSON エクスポート後 DELETE", style="List Bullet")
    doc.add_paragraph("(c) gps_points: 365日超のGPSデータを S3 に JSON エクスポート後 DELETE", style="List Bullet")
    doc.add_paragraph("S3キー: s3://pms-uploads-tiramis/archives/{テーブル名}/{YYYY-MM-DD}/", style="List Bullet")
    doc.add_paragraph("推定実行時間: 1秒〜数分（蓄積データ量に依存）", style="List Bullet")
    doc.add_paragraph("DB負荷: 高（大量DELETE + S3書き込み。gps_pointsは年間5,000万行規模）", style="List Bullet")

    # --- 5. スケジュール分散の推奨 ---
    doc.add_heading("5. スケジュール分散の推奨（現状の課題）", level=1)
    doc.add_paragraph(
        "現在、全7ジョブが 15:00 UTC（00:00 JST）に集中している。"
        "特にハウスキープ（重い DELETE + S3書き込み）と他ジョブが同時実行されると、"
        "DB接続プールを圧迫しレスポンス遅延の原因となる。"
    )
    doc.add_paragraph("■ 推奨スケジュール（分散案）")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["ジョブ", "推奨 CRON (UTC)", "JST", "理由"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    recommended = [
        ("generate-slots", "0 16 * * *", "01:00", "業務時間外・低負荷"),
        ("generate-tasks", "0 15 * * *", "00:00", "日付変更直後に当日タスクを生成"),
        ("generate-training-slots", "30 16 * * *", "01:30", "面接スロット生成後"),
        ("evaluate-distributors", "0 19 * * 1", "04:00 月曜", "週初め早朝・重め処理"),
        ("check-visa-expiry", "0 18 * * *", "03:00", "業務開始前にアラート準備"),
        ("check-alert-definitions", "0 21 * * *", "06:00", "始業直前にアラート確認可能"),
        ("housekeep", "0 17 * * *", "02:00", "最も重い処理を単独実行"),
    ]
    for r in recommended:
        add_table_row(table, r)
    set_table_style(table)

    # --- 6. メンテナンス窓 ---
    doc.add_heading("6. メンテナンス窓", level=1)
    doc.add_paragraph(
        "CRON ジョブが実行されず、かつユーザーのアクセスが想定されない時間帯を "
        "メンテナンス窓として定義する。RDS スケールアップ・マイグレーション等に使用する。"
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["区分", "UTC", "JST"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    add_table_row(table, ["推奨メンテナンス窓（現行スケジュール）", "17:00 - 20:00", "02:00 - 05:00"])
    add_table_row(table, ["推奨メンテナンス窓（分散案適用後）", "22:00 - 14:00", "07:00 - 23:00 の間で調整"])
    add_table_row(table, ["禁止時間帯（業務ピーク）", "00:00 - 09:00", "09:00 - 18:00"])
    set_table_style(table)

    # --- 7. 障害時対応 ---
    doc.add_heading("7. 障害時対応", level=1)
    doc.add_paragraph("■ ジョブ失敗時")
    doc.add_paragraph("ログファイル（/tmp/pms-cron-*.log）を確認し、HTTP ステータスとエラーメッセージを特定する。", style="List Bullet")
    doc.add_paragraph("手動再実行: curl -H 'Authorization: Bearer $CRON_SECRET' http://localhost:3000/api/cron/{ジョブ名}", style="List Bullet")
    doc.add_paragraph("DB接続エラーの場合は RDS の状態・接続数上限を確認する。", style="List Bullet")

    doc.add_paragraph("■ 二重実行防止")
    doc.add_paragraph("各ジョブは冪等性を持つ設計（同日・同条件の重複チェック）。"
                      "万一二重実行されてもデータ不整合は発生しない。", style="List Bullet")

    # --- 8. 監視 ---
    doc.add_heading("8. 監視", level=1)
    doc.add_paragraph("現状: ログファイルベースの手動確認")
    doc.add_paragraph("■ 将来的な監視強化案")
    doc.add_paragraph("ジョブ成功/失敗を SystemSetting に記録し、ダッシュボードに表示", style="List Bullet")
    doc.add_paragraph("失敗時に管理者通知（admin_notifications）を生成", style="List Bullet")
    doc.add_paragraph("CloudWatch Logs への出力切替（ログローテーション不要化）", style="List Bullet")

    filepath = os.path.join(OUTPUT_DIR, "PMS_ジョブスケジュール設計書.docx")
    doc.save(filepath)
    print(f"生成完了: {filepath}")


# ============================================================
# 2. SLA設計書
# ============================================================
def generate_sla_doc():
    doc = docx.Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Yu Gothic"
    font.size = Pt(10)

    add_title_page(doc, "PMS SLA 定義", "可用性・パフォーマンス設計書")

    # --- 1. 概要 ---
    doc.add_heading("1. 概要", level=1)
    doc.add_paragraph(
        "本書では、PMS（ポスティング管理システム）およびECポータルの "
        "サービスレベル目標（SLO）・可用性設計・パフォーマンス基準・キャパシティ計画を定義する。"
    )

    # --- 2. システム構成 ---
    doc.add_heading("2. システム構成（2026年3月時点）", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["レイヤー", "サービス", "スペック", "冗長性"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    rows = [
        ("CDN", "CloudFront", "d1mzs3dojvfqpz.cloudfront.net", "グローバル分散"),
        ("LB", "ALB", "50/50 トラフィック分散", "2 AZ 冗長"),
        ("App", "EC2 × 2台", "t3.large (2vCPU / 8GB) × 2", "Multi-AZ"),
        ("App", "PM2 Cluster", "2プロセス/台 = 合計4プロセス", "プロセス自動復旧"),
        ("DB", "RDS MySQL", "db.t3.small (2vCPU / 2GB)", "Single-AZ ※要検討"),
        ("Storage", "S3", "pms-uploads-tiramis", "99.999999999% 耐久性"),
    ]
    for r in rows:
        add_table_row(table, r)
    set_table_style(table)

    p = doc.add_paragraph()
    run = p.add_run("※ RDS は 2026年3月4日に db.t3.micro → db.t3.small にスケールアップ予定")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)

    # --- 3. SLO 定義 ---
    doc.add_heading("3. サービスレベル目標（SLO）", level=1)

    doc.add_heading("3.1 可用性", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["対象", "SLO", "許容ダウンタイム/月", "備考"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    add_table_row(table, ["管理画面 (PMS)", "99.5%", "約3.6時間/月", "業務時間帯 (9:00-18:00 JST) 優先"])
    add_table_row(table, ["ECポータル", "99.5%", "約3.6時間/月", "24時間アクセス想定"])
    add_table_row(table, ["モバイルアプリAPI", "99.5%", "約3.6時間/月", "配布時間帯 (6:00-20:00 JST) 重要"])
    add_table_row(table, ["GPS APIエンドポイント", "99.9%", "約43分/月", "高頻度（10秒間隔）・データ欠損影響大"])
    add_table_row(table, ["CRONジョブ", "99%", "約7.3時間/月", "冪等設計のため翌日自動リカバリ可能"])
    set_table_style(table)

    doc.add_heading("3.2 レスポンスタイム", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["対象", "P50", "P95", "P99"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    add_table_row(table, ["管理画面ページ遷移（SSR）", "< 500ms", "< 1.5s", "< 3s"])
    add_table_row(table, ["API レスポンス（一般）", "< 200ms", "< 500ms", "< 1s"])
    add_table_row(table, ["GPS 座標受信 API", "< 50ms", "< 100ms", "< 200ms"])
    add_table_row(table, ["軌跡データ取得（最大2,880点）", "< 500ms", "< 1.5s", "< 3s"])
    add_table_row(table, ["ECポータルページ表示", "< 800ms", "< 2s", "< 4s"])
    add_table_row(table, ["S3 ファイルアップロード", "< 2s", "< 5s", "< 10s"])
    set_table_style(table)

    # --- 4. キャパシティ計画 ---
    doc.add_heading("4. キャパシティ計画", level=1)

    doc.add_heading("4.1 同時接続数", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["シナリオ", "同時接続数", "推定リクエスト"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    add_table_row(table, ["通常運用", "30〜50人", "~5 req/秒"])
    add_table_row(table, ["ピーク時", "100人", "~10 req/秒"])
    add_table_row(table, ["最大想定", "200人", "~20 req/秒"])
    set_table_style(table)

    doc.add_paragraph(
        "Node.js (PM2 cluster × 4プロセス) の処理能力は I/O バウンドで数百 req/秒。"
        "現行スペックで 200 人同時接続まで対応可能。"
    )

    doc.add_heading("4.2 DB 接続プール", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["項目", "db.t3.micro (旧)", "db.t3.small (現行)"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    add_table_row(table, ["max_connections", "60", "150"])
    add_table_row(table, ["Prisma プール/プロセス", "5", "5"])
    add_table_row(table, ["プロセス合計 (4台)", "20", "20"])
    add_table_row(table, ["CRON + 管理ツール予備", "~10", "~10"])
    add_table_row(table, ["ピーク時想定使用", "~50", "~50"])
    add_table_row(table, ["余裕率", "17% (危険)", "67% (安全)"])
    set_table_style(table)

    doc.add_heading("4.3 データ増加予測", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["テーブル", "増加ペース", "年間蓄積", "対策"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    add_table_row(table, ["gps_points", "~144,000行/日", "~5,000万行/年", "365日超をS3アーカイブ (housekeep)"])
    add_table_row(table, ["audit_logs", "~500行/日", "~18万行/年", "90日超をS3アーカイブ (housekeep)"])
    add_table_row(table, ["admin_notifications", "~100行/日", "~3.6万行/年", "30日超を削除 (housekeep)"])
    add_table_row(table, ["progress_events", "~500行/日", "~18万行/年", "アーカイブ未実装（要検討）"])
    set_table_style(table)

    # --- 5. 可用性設計 ---
    doc.add_heading("5. 可用性設計", level=1)

    doc.add_heading("5.1 現状の冗長構成", level=2)
    doc.add_paragraph("ALB + EC2 × 2台（ap-northeast-1a / 1c）により、1台障害時もサービス継続可能。", style="List Bullet")
    doc.add_paragraph("PM2 cluster mode（2プロセス/台）により、1プロセスクラッシュ時も自動復旧。", style="List Bullet")
    doc.add_paragraph("CloudFront による静的アセット配信で、EC2 障害時も CSS/JS は表示可能。", style="List Bullet")
    doc.add_paragraph("CRONジョブは両EC2で登録されているが、localhost:3000 呼び出しのため各EC2で独立実行。"
                      "冪等設計により二重実行しても問題なし。", style="List Bullet")

    doc.add_heading("5.2 単一障害点（SPOF）", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["コンポーネント", "SPOF?", "影響", "対策案"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    add_table_row(table, ["RDS MySQL", "はい", "全サービス停止", "Multi-AZ 有効化（推奨）"])
    add_table_row(table, ["ALB", "いいえ", "-", "AWS マネージド冗長"])
    add_table_row(table, ["EC2 (1台)", "いいえ", "50% 処理能力低下", "ALB ヘルスチェックで切離し"])
    add_table_row(table, ["S3", "いいえ", "-", "AWS マネージド（11ナインの耐久性）"])
    add_table_row(table, ["CloudFront", "いいえ", "-", "グローバル分散"])
    add_table_row(table, ["SMTP (Gmail)", "はい", "メール送信不可", "SES 移行を検討"])
    set_table_style(table)

    p = doc.add_paragraph()
    run = p.add_run("★ RDS が最大の SPOF。Multi-AZ 化により自動フェイルオーバー（60〜120秒）が可能になる。")
    run.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
    run.bold = True

    # --- 6. パフォーマンスボトルネック分析 ---
    doc.add_heading("6. パフォーマンスボトルネック分析", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["優先度", "ボトルネック", "影響", "対策"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    add_table_row(table, ["高", "gps_points テーブル肥大化", "軌跡クエリ遅延・ストレージ圧迫",
                          "パーティショニング or TTLインデックス"])
    add_table_row(table, ["高", "通知ポーリング (30秒間隔)", "常時リクエスト負荷",
                          "SSE/WebSocket への移行"])
    add_table_row(table, ["中", "CRON 同時実行", "DB接続プール圧迫・レスポンス遅延",
                          "スケジュール分散（ジョブ設計書参照）"])
    add_table_row(table, ["中", "RDS Single-AZ", "障害時に全サービス停止",
                          "Multi-AZ 有効化"])
    add_table_row(table, ["低", "軌跡ビューア（2,880点一括取得）", "大量データ転送",
                          "ページネーション or サマリー表示"])
    set_table_style(table)

    # --- 7. スケールアップ計画 ---
    doc.add_heading("7. スケールアップ計画", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["フェーズ", "条件（トリガー）", "アクション", "コスト増", "期待効果"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    add_table_row(table, ["Phase 1\n（現行）", "-", "db.t3.micro → db.t3.small\nPM2 fork → cluster",
                          "+$16/月", "接続数 60→150\nCPU活用率 2倍"])
    add_table_row(table, ["Phase 2", "配布員 100人超\nor 接続数 100超",
                          "RDS Multi-AZ 有効化", "+$31/月", "DB SPOF 解消\n自動フェイルオーバー"])
    add_table_row(table, ["Phase 3", "同時接続 200人超\nor レスポンス P95 > 2s",
                          "EC2 t3.large → t3.xlarge\n(4vCPU/16GB)",
                          "+$60/月/台", "処理能力 2倍\nPM2 4プロセス/台"])
    add_table_row(table, ["Phase 4", "gps_points 1億行超", "RDS db.t3.small → db.t3.medium\nリードレプリカ追加",
                          "+$62/月", "クエリ性能向上\n読み取り負荷分散"])
    set_table_style(table)

    # --- 8. メンテナンス窓 ---
    doc.add_heading("8. メンテナンス窓", level=1)
    doc.add_paragraph(
        "計画メンテナンスは以下の時間帯に実施する。ユーザーへの事前告知はメール + 管理画面通知で行う。"
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["種別", "推奨時間帯 (JST)", "所要時間目安"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    add_table_row(table, ["RDS スケールアップ", "02:00 - 05:00", "5〜10分"])
    add_table_row(table, ["RDS Multi-AZ 切替", "02:00 - 05:00", "10〜20分"])
    add_table_row(table, ["EC2 インスタンスタイプ変更", "02:00 - 05:00", "5〜10分（1台ずつローリング）"])
    add_table_row(table, ["DB マイグレーション（スキーマ変更）", "02:00 - 05:00", "1〜5分"])
    add_table_row(table, ["通常デプロイ", "随時（AllAtOnce）", "3〜5分（自動）"])
    set_table_style(table)

    # --- 9. 障害対応フロー ---
    doc.add_heading("9. 障害対応フロー", level=1)

    doc.add_heading("9.1 EC2 障害", level=2)
    doc.add_paragraph("1. ALB ヘルスチェックが異常を検知 → 自動的にトラフィック切離し")
    doc.add_paragraph("2. もう1台のEC2で全トラフィックを処理（縮退運転）")
    doc.add_paragraph("3. SSH で障害EC2を確認: pm2 status, pm2 logs pms, free -h")
    doc.add_paragraph("4. 必要に応じて pm2 restart pms or EC2 再起動")

    doc.add_heading("9.2 RDS 障害", level=2)
    doc.add_paragraph("1. アプリから DB 接続エラー発生 → 全サービス停止")
    doc.add_paragraph("2. RDS コンソールで状態確認")
    doc.add_paragraph("3. Multi-AZ の場合: 自動フェイルオーバー（60〜120秒）")
    doc.add_paragraph("4. Single-AZ の場合: 手動復旧待ち（数分〜数十分）")
    doc.add_paragraph("5. 復旧後、PM2 が自動再接続 → サービス復旧")

    doc.add_heading("9.3 デプロイ障害", level=2)
    doc.add_paragraph("1. GitHub Actions or CodeDeploy が失敗")
    doc.add_paragraph("2. CodeDeploy コンソールでエラー内容確認")
    doc.add_paragraph("3. 前回デプロイが動作中なら影響なし（新バージョンが適用されないだけ）")
    doc.add_paragraph("4. SSH で手動確認: pm2 status, /tmp/pms-cron-*.log")

    # --- 10. 改版履歴 ---
    doc.add_heading("10. 改版履歴", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["日付", "版", "変更内容", "担当"]):
        hdr.cells[i].text = h
    style_header_row(hdr)
    add_table_row(table, ["2026-03-04", "1.0", "初版作成", "Claude / Kim"])
    set_table_style(table)

    filepath = os.path.join(OUTPUT_DIR, "PMS_SLA設計書.docx")
    doc.save(filepath)
    print(f"生成完了: {filepath}")


if __name__ == "__main__":
    generate_job_schedule_doc()
    generate_sla_doc()
    print("\n全ドキュメント生成完了")
