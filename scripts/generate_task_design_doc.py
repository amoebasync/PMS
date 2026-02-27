"""
CRM タスク機能改修（定期タスク自動生成 & 柔軟な担当者アサイン）設計書生成スクリプト
"""
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_PATH = os.path.expanduser("~/Downloads/PMS_設計書/PMS_CRMタスク機能改修_設計書.docx")

doc = docx.Document()

# ===== スタイル設定 =====
style = doc.styles['Normal']
font = style.font
font.name = 'Yu Gothic'
font.size = Pt(10.5)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Yu Gothic'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

# ===== 表紙 =====
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PMS Pro\nCRMタスク機能改修 設計書')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('株式会社ティラミス 社内向け基幹管理システム')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(4):
    doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('作成日: 2026-02-27\nバージョン: 1.0\n\n定期タスク自動生成（リマインド/ToDo）\n柔軟な担当者アサイン（個人・部署・支店）')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_page_break()

# ===== ヘルパー関数 =====
def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# ===== 1. 概要 =====
doc.add_heading('1. 概要', level=1)
doc.add_paragraph(
    '本設計書は、PMS Pro の CRM タスク機能を改修し、「定期タスクの自動生成（リマインド/ToDo機能）」'
    'および「柔軟な担当者アサイン（個人・部署・支店単位）」を追加した設計内容を記述する。'
)

doc.add_heading('1.1 目的', level=2)
bullet('タスクをカテゴリ（営業/現場/アドミン）で分類し、業務に応じた関連先を紐付け')
bullet('定期タスクテンプレートを作成し、サイクル（毎日/毎週/毎月/毎年）に基づいてタスクを自動生成')
bullet('担当者を社員個人だけでなく、部署・支店単位でも柔軟にアサイン可能に')
bullet('共有タスク（SHARED: 1タスクを複数人で共有）と個別タスク（INDIVIDUAL: 人数分のタスクを個別生成）の使い分け')
bullet('マイタスク機能により、自分にアサインされたタスク（個人・所属部署・所属支店）を素早くフィルタ')
bullet('CRONジョブによるタスク自動生成（毎日00:00実行）')

doc.add_heading('1.2 スコープ', level=2)
bullet('対象ユーザー：社内管理者（pms_session Cookie 認証）')
bullet('画面：/crm/tasks（タスク一覧タブ + 定期タスク設定タブ）')
bullet('API：担当者検索1 + テンプレートCRUD 2 + タスクCRUD拡張 2 + CRON 1 = 計6エンドポイント')
bullet('CRON：定期タスク自動生成バッチ（Bearer トークン認証）')
bullet('監査ログ：テンプレートCRUD・CRON生成の全操作を記録')

# ===== 2. データベース設計 =====
doc.add_heading('2. データベース設計', level=1)

doc.add_heading('2.1 Enum 定義', level=2)
doc.add_paragraph('TaskCategory（タスクカテゴリ）')
add_table(['値', '説明'], [
    ['SALES', '営業タスク（顧客紐付け）'],
    ['FIELD', '現場タスク（支店紐付け）'],
    ['ADMIN', 'アドミンタスク（紐付けなし）'],
])

doc.add_paragraph('')
doc.add_paragraph('RecurrenceType（繰り返しタイプ）')
add_table(['値', '説明', 'recurrenceValue 例'], [
    ['ONCE', '一回のみ', '(なし)'],
    ['DAILY', '毎日', '(なし)'],
    ['WEEKLY', '毎週（曜日指定）', '"1,3,5"（月水金） ※0=日,1=月,...,6=土'],
    ['MONTHLY', '毎月（日付指定）', '"15"（15日）'],
    ['YEARLY', '毎年（月日指定）', '"03-15"（3月15日）'],
])

doc.add_paragraph('')
doc.add_paragraph('TaskCompletionRule（完了条件）')
add_table(['値', '説明'], [
    ['SHARED', '1つのタスクを複数人で共有。誰か1人が完了すればOK'],
    ['INDIVIDUAL', '担当者ごとに独立タスクを生成。各自が個別に完了させる'],
])

doc.add_heading('2.2 Task テーブル拡張', level=2)
doc.add_paragraph('既存の tasks テーブルに以下のカラムを追加:')
add_table(['カラム名', '型', '説明'], [
    ['category', 'TaskCategory?', 'タスクカテゴリ（SALES/FIELD/ADMIN）'],
    ['branch_id', 'Int? (FK)', '支店への外部キー（現場タスク用）'],
    ['schedule_id', 'Int? (FK)', '配送スケジュールへの外部キー'],
    ['template_id', 'Int? (FK)', '生成元テンプレートへの外部キー'],
])

doc.add_heading('2.3 TaskAssignee テーブル（新規）', level=2)
doc.add_paragraph('テーブル名: task_assignees — タスクと担当者の中間テーブル。社員・部署・支店を横断して紐付け可能。')
add_table(['カラム名', '型', '説明'], [
    ['id', 'Int (PK)', '自動採番'],
    ['task_id', 'Int (FK)', 'タスクID（Cascade削除）'],
    ['employee_id', 'Int? (FK)', '社員ID'],
    ['department_id', 'Int? (FK)', '部署ID'],
    ['branch_id', 'Int? (FK)', '支店ID'],
])
doc.add_paragraph('UNIQUE制約: (task_id, employee_id, department_id, branch_id)')

doc.add_heading('2.4 TaskTemplate テーブル（新規）', level=2)
doc.add_paragraph('テーブル名: task_templates — 定期タスクのテンプレート定義マスタ。')
add_table(['カラム名', '型', '説明'], [
    ['id', 'Int (PK)', '自動採番'],
    ['title', 'VarChar(200)', 'タスクタイトル'],
    ['description', 'Text?', '説明'],
    ['category', 'TaskCategory', 'タスクカテゴリ（必須）'],
    ['priority', 'VarChar(10)', '優先度（HIGH/MEDIUM/LOW）。デフォルト: MEDIUM'],
    ['completion_rule', 'TaskCompletionRule', '完了条件。デフォルト: SHARED'],
    ['customer_id', 'Int? (FK)', '顧客紐付け（営業タスク用）'],
    ['branch_id', 'Int? (FK)', '支店紐付け（現場タスク用）'],
    ['schedule_id', 'Int? (FK)', '配送スケジュール紐付け'],
    ['recurrence_type', 'RecurrenceType', '繰り返しタイプ。デフォルト: ONCE'],
    ['recurrence_value', 'VarChar(100)?', '繰り返し詳細値（曜日/日付/月日）'],
    ['target_employee_ids', 'Json?', '対象社員IDの配列 例: [1, 5, 10]'],
    ['target_department_ids', 'Json?', '対象部署IDの配列 例: [2, 3]'],
    ['target_branch_ids', 'Json?', '対象支店IDの配列 例: [1]'],
    ['is_active', 'Boolean', '有効/無効フラグ。デフォルト: true'],
    ['last_generated_at', 'DateTime?', '最終タスク生成日時'],
    ['created_at', 'DateTime', '作成日時'],
    ['updated_at', 'DateTime', '更新日時'],
])

doc.add_heading('2.5 逆リレーション追加', level=2)
add_table(['モデル', '追加リレーション'], [
    ['Employee', 'taskAssignments: TaskAssignee[]'],
    ['Department', 'taskAssignees: TaskAssignee[]'],
    ['Branch', 'tasksByBranch: Task[], taskAssignees: TaskAssignee[], taskTemplates: TaskTemplate[]'],
    ['Customer', 'taskTemplates: TaskTemplate[]'],
    ['DistributionSchedule', 'tasks: Task[], taskTemplates: TaskTemplate[]'],
])

# ===== 3. API 設計 =====
doc.add_heading('3. API 設計', level=1)

doc.add_heading('3.1 担当者横断検索 API', level=2)
doc.add_paragraph('ファイル: src/app/api/staff/search-assignees/route.ts（新規）')
add_table(['メソッド', 'パス', '認証', '説明'], [
    ['GET', '/api/staff/search-assignees?q=keyword', 'pms_session', '社員・部署・支店を横断検索'],
])
doc.add_paragraph('レスポンス例:')
doc.add_paragraph(
    '[\n'
    '  { "type": "employee", "id": 1, "label": "金 建希", "sub": "営業部" },\n'
    '  { "type": "department", "id": 2, "label": "営業部", "sub": "社員5名" },\n'
    '  { "type": "branch", "id": 1, "label": "高田馬場", "sub": "支店（社員3名）" }\n'
    ']'
)

doc.add_heading('3.2 テンプレート CRUD API', level=2)
doc.add_paragraph('ファイル: src/app/api/task-templates/route.ts, [id]/route.ts（新規）')
add_table(['メソッド', 'パス', '認証', '説明'], [
    ['GET', '/api/task-templates', 'pms_session', 'テンプレート一覧（category, isActive フィルタ対応）'],
    ['POST', '/api/task-templates', 'pms_session', 'テンプレート作成（トランザクション + 監査ログ）'],
    ['GET', '/api/task-templates/[id]', 'pms_session', 'テンプレート詳細取得'],
    ['PUT', '/api/task-templates/[id]', 'pms_session', 'テンプレート更新（トランザクション + 監査ログ）'],
    ['DELETE', '/api/task-templates/[id]', 'pms_session', 'テンプレート削除（トランザクション + 監査ログ）'],
])

doc.add_heading('3.3 タスク API 拡張', level=2)
doc.add_paragraph('ファイル: src/app/api/tasks/route.ts, [id]/route.ts（既存修正）')
doc.add_paragraph('GET 拡張:')
bullet('category フィルタ追加')
bullet('myTasks=true パラメータ: ログインユーザーの個人ID・所属部署ID・所属支店IDで assignees 中間テーブルを検索')
bullet('include に branch, schedule, template, assignees を追加')
doc.add_paragraph('POST / PUT 拡張:')
bullet('category, branchId, scheduleId, assignees パラメータ受付')
bullet('assignees はトランザクション内で TaskAssignee を一括作成/差し替え')

doc.add_heading('3.4 定期タスク自動生成 CRON API', level=2)
doc.add_paragraph('ファイル: src/app/api/cron/generate-tasks/route.ts（新規）')
add_table(['メソッド', 'パス', '認証', '説明'], [
    ['GET', '/api/cron/generate-tasks', 'Bearer CRON_SECRET', '有効なテンプレートから本日分のタスクを自動生成'],
])
doc.add_paragraph('処理フロー:')
bullet('1. isActive: true のテンプレートを全取得')
bullet('2. 各テンプレートの recurrenceType + recurrenceValue と本日を照合（shouldGenerateToday 関数）')
bullet('3. 同日に同テンプレートから既に生成済みかチェック（重複防止）')
bullet('4. completionRule に従いタスク生成:')
bullet('   SHARED: Task 1件 + TaskAssignee に全担当者を紐付け')
bullet('   INDIVIDUAL: 対象社員を展開（部署→所属社員、支店→所属社員）→ 人数分の Task を個別作成')
bullet('5. dueDate = 本日、status = PENDING')
bullet('6. lastGeneratedAt を更新')
bullet('7. 監査ログ（actorType: SYSTEM）')

# ===== 4. ミドルウェア / インフラ =====
doc.add_heading('4. ミドルウェア・インフラ設定', level=1)

doc.add_heading('4.1 ミドルウェア', level=2)
doc.add_paragraph('ファイル: src/middleware.ts')
bullet('/api/cron/generate-tasks を公開パス（認証スキップ）に追加')
bullet('CRON APIは独自の Bearer トークン認証を行うため、ミドルウェアの pms_session チェックは不要')

doc.add_heading('4.2 CRON 登録（CodeDeploy）', level=2)
doc.add_paragraph('ファイル: scripts/after_install.sh')
bullet('デプロイ時に crontab に自動登録（重複チェック付き）')
bullet('面接スロット生成: 毎日01:00 → /api/cron/generate-slots')
bullet('定期タスク生成: 毎日00:00 → /api/cron/generate-tasks')
bullet('両方とも CRON_SECRET 環境変数を使用した Bearer 認証')
bullet('ALBのURL（https://pms.tiramis.co.jp）にリクエスト送信')

# ===== 5. フロントエンド設計 =====
doc.add_heading('5. フロントエンド設計', level=1)
doc.add_paragraph('ファイル: src/app/crm/tasks/page.tsx（既存大幅改修）')

doc.add_heading('5.1 タブ構成', level=2)
doc.add_paragraph('既存のタスク一覧を「タスク一覧」タブとし、新たに「定期タスク設定」タブを追加。')
add_table(['タブ名', '内容'], [
    ['タスク一覧', 'タスクの一覧表示・フィルタ・CRUD'],
    ['定期タスク設定', 'テンプレートの一覧・作成・編集・削除・有効/無効切替'],
])

doc.add_heading('5.2 タスク一覧タブの改修内容', level=2)
bullet('カテゴリバッジ表示（営業=blue, 現場=green, アドミン=slate）')
bullet('カテゴリフィルタ追加（ドロップダウン: すべて/営業/現場/アドミン）')
bullet('マイタスクフィルタ（トグルボタン: 自分にアサインされたタスクのみ表示）')
bullet('担当者列にマルチ担当者表示（複数人の場合は「xxx 他N名」）')
bullet('関連先列にカテゴリ別の情報表示（顧客名/支店名/配布員名）')

doc.add_heading('5.3 タスク作成/編集モーダル', level=2)
bullet('タスク種類ボタン（指定なし/営業/現場/アドミン）追加')
bullet('種類に応じた関連先UIの動的切替:')
bullet('  営業 → 顧客オートコンプリート')
bullet('  現場 → 支店ドロップダウン')
bullet('  アドミン → 紐付けなし')
bullet('  指定なし → 顧客 or 配布員オートコンプリート（後方互換）')
bullet('担当者マルチセレクト（search-assignees API 連携）')

doc.add_heading('5.4 定期タスク設定タブ', level=2)
doc.add_paragraph('テンプレート一覧テーブル:')
bullet('タイトル、種類バッジ、サイクル表示、優先度、担当者数、完了条件、有効/無効トグル、最終生成日時')
doc.add_paragraph('テンプレート作成/編集モーダル:')
bullet('タイトル・説明入力')
bullet('タスク種類（SALES/FIELD/ADMIN）→ 関連先UIの出し分け')
bullet('サイクル設定:')
bullet('  繰り返しタイプ ドロップダウン')
bullet('  WEEKLY → 曜日マルチ選択ボタン（日〜土、日=赤、土=青）')
bullet('  MONTHLY → 日付ドロップダウン（1〜31日）')
bullet('  YEARLY → 月 + 日 ドロップダウン')
bullet('担当者オートコンプリートUI（社員/部署/支店を区別するアイコン+ラベル、タグ表示）')
bullet('担当者2名以上の場合 → 完了条件ラジオ（共有タスク / 個別タスク）')
bullet('有効/無効トグル')

doc.add_heading('5.5 AssigneeMultiSelect コンポーネント', level=2)
doc.add_paragraph('担当者のマルチセレクトUIコンポーネント。以下の機能を持つ:')
bullet('テキスト入力 → /api/staff/search-assignees でリアルタイム検索')
bullet('検索結果にタイプ別アイコン（社員=person、部署=people、支店=building）を表示')
bullet('選択済み項目はタイプ別色分けタグで表示（社員=blue、部署=purple、支店=amber）')
bullet('×ボタンで個別削除可能')
bullet('既に選択済みの項目は検索結果から自動除外')

# ===== 6. ファイル一覧 =====
doc.add_heading('6. ファイル一覧', level=1)

doc.add_heading('6.1 新規作成ファイル（4件）', level=2)
add_table(['ファイルパス', '概要'], [
    ['src/app/api/staff/search-assignees/route.ts', '担当者横断検索API'],
    ['src/app/api/task-templates/route.ts', 'テンプレート一覧・作成（GET/POST）'],
    ['src/app/api/task-templates/[id]/route.ts', 'テンプレート詳細・更新・削除（GET/PUT/DELETE）'],
    ['src/app/api/cron/generate-tasks/route.ts', '定期タスク自動生成CRON'],
])

doc.add_heading('6.2 修正ファイル（5件）', level=2)
add_table(['ファイルパス', '修正内容'], [
    ['prisma/schema.prisma', 'Enum 3つ + Model 2つ + Task拡張 + 逆リレーション追加'],
    ['src/app/api/tasks/route.ts', 'GET/POST に category, assignees, myTasks 対応追加'],
    ['src/app/api/tasks/[id]/route.ts', 'PUT に category, assignees 対応追加'],
    ['src/app/crm/tasks/page.tsx', 'タブ追加、カテゴリUI、マイタスク、マルチ担当者、定期タスク設定'],
    ['src/middleware.ts', '/api/cron/generate-tasks を公開パスに追加'],
    ['scripts/after_install.sh', 'generate-tasks CRON登録を追加'],
])

# ===== 7. 検証方法 =====
doc.add_heading('7. 検証方法', level=1)
bullet('npx prisma db push でスキーマ反映を確認')
bullet('npm run dev でサーバー起動、/crm/tasks 画面にアクセス')
bullet('タブ切替（タスク一覧 ↔ 定期タスク設定）が正常に動作すること')
bullet('タスク作成モーダル: カテゴリ選択で関連先UIが動的に切替わること')
bullet('担当者マルチセレクト: 社員・部署・支店を検索して追加・削除できること')
bullet('テンプレート作成: WEEKLY選択で曜日ボタン、MONTHLY選択で日付ドロップダウンが表示されること')
bullet('テンプレート有効/無効トグルが正常に動作すること')
bullet('CRON手動実行: curl -H "Authorization: Bearer $CRON_SECRET" http://localhost:3000/api/cron/generate-tasks')
bullet('CRON実行後、テンプレートに基づくタスクがタスク一覧に表示されること')
bullet('SHARED: 1タスクに複数担当者が紐付くこと / INDIVIDUAL: 人数分の個別タスクが生成されること')
bullet('マイタスクフィルタ: 自分の個人・部署・支店に関連するタスクのみ表示されること')
bullet('同日に同テンプレートからの二重生成が防止されること')

# 保存
doc.save(OUTPUT_PATH)
print(f"設計書を生成しました: {OUTPUT_PATH}")
