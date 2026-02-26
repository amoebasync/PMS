#!/usr/bin/env python3
"""
応募者管理機能 設計書 & 機能一覧更新 生成スクリプト
"""
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from copy import deepcopy
import os

OUTPUT_DIR = os.path.expanduser("~/Downloads/PMS_設計書")

# ====================================================================
# ユーティリティ
# ====================================================================

def set_cell_shading(cell, color_hex):
    """セルの背景色を設定"""
    shading_elm = docx.oxml.parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """ヘッダー付きテーブルを追加"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1E293B")

    # データ行
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F8FAFC")

    return table


# ====================================================================
# 1. 設計書
# ====================================================================

def create_design_doc():
    doc = Document()

    # デフォルトフォント設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic UI'
    font.size = Pt(10)

    # --- 表紙 ---
    for _ in range(4):
        doc.add_paragraph('')

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('PMS Pro\n応募者管理・面接日程調整機能 設計書')
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('株式会社ティラミス 社内向け基幹管理システム')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph('')

    # 文書情報テーブル
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('文書バージョン', '1.0'),
        ('作成日', '2026年02月27日'),
        ('対象システム', 'PMS Pro（ポスティング管理システム）'),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        for cell in info_table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_page_break()

    # === 1. 概要 ===
    doc.add_heading('1. 概要', level=1)
    doc.add_paragraph(
        '本設計書は、PMS Pro に新規追加した「応募者管理・面接日程調整機能」の設計内容を記述する。'
        '外部ツール（TimeRex 等）を使用せず、システム内で面接枠の作成から予約まで完結させるネイティブ実装である。'
    )

    doc.add_heading('1.1 目的', level=2)
    purposes = [
        '応募者情報（氏名・連絡先・国籍・在留資格等）の一元管理',
        '管理者がカレンダーUIで面接スロット（空き枠）を作成・管理',
        '応募者が公開フォームから面接枠を自ら選択・予約（競合防止のトランザクション処理）',
        '面接実施後の評価入力（語学力・コミュニケーション・印象等 1-5段階）',
        '応募フロー（面接待ち→研修待ち→研修完了）および採用ステータス（選考中→採用→不採用）の管理',
        '応募受付確認メール・採用通知メールの自動送信（日本語/英語対応）',
    ]
    for p in purposes:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('1.2 スコープ', level=2)
    scopes = [
        '対象ユーザー：応募者（認証なし・公開フォーム） / 社内管理者（pms_session Cookie）',
        '画面：応募フォーム（/apply）、管理者カレンダー＆一覧画面（/applicants）',
        'API：公開5エンドポイント + 管理者5エンドポイント',
        'メール：面接確認通知（ja/en）、採用通知（ja/en）',
        '監査ログ：応募・評価更新・ステータス変更・スロットCRUDの全操作を記録',
    ]
    for s in scopes:
        doc.add_paragraph(s, style='List Bullet')

    # === 2. データベース設計 ===
    doc.add_heading('2. データベース設計', level=1)

    # 2.1 Enum
    doc.add_heading('2.1 Enum 定義', level=2)

    doc.add_paragraph('ApplicantFlowStatus（応募フロー状態）')
    add_styled_table(doc,
        ['値', '日本語ラベル', '説明'],
        [
            ['INTERVIEW_WAITING', '面接待ち', '応募直後のデフォルト状態。面接実施を待機中'],
            ['TRAINING_WAITING', '研修待ち', '面接通過後、研修スケジュール待ち'],
            ['TRAINING_COMPLETED', '研修完了', '研修完了。配属可能な状態'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph('ApplicantHiringStatus（採用ステータス）')
    add_styled_table(doc,
        ['値', '日本語ラベル', '説明'],
        [
            ['IN_PROGRESS', '選考中', '選考プロセス進行中（デフォルト）'],
            ['HIRED', '採用', '採用決定。ステータス変更時に採用通知メール自動送信'],
            ['REJECTED', '不採用', '不採用'],
        ]
    )

    # 2.2 JobCategory
    doc.add_heading('2.2 JobCategory テーブル（応募職種マスタ）', level=2)
    doc.add_paragraph('テーブル名: job_categories')
    add_styled_table(doc,
        ['カラム名（DB）', 'Prismaフィールド名', 'データ型', 'NULL許可', '説明'],
        [
            ['id', 'id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['name_ja', 'nameJa', 'VARCHAR(100)', '不可', '職種名（日本語）'],
            ['name_en', 'nameEn', 'VARCHAR(100)', '可', '職種名（英語）'],
            ['is_active', 'isActive', 'BOOLEAN', '不可', '有効フラグ（default: true）'],
            ['sort_order', 'sortOrder', 'INT', '不可', '表示順（default: 100）'],
            ['created_at', 'createdAt', 'DATETIME', '不可', '作成日時'],
            ['updated_at', 'updatedAt', 'DATETIME', '不可', '更新日時'],
        ]
    )

    # 2.3 InterviewSlot
    doc.add_heading('2.3 InterviewSlot テーブル（面接スロット）', level=2)
    doc.add_paragraph('テーブル名: interview_slots')
    add_styled_table(doc,
        ['カラム名（DB）', 'Prismaフィールド名', 'データ型', 'NULL許可', '説明'],
        [
            ['id', 'id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['start_time', 'startTime', 'DATETIME', '不可', '面接開始日時'],
            ['end_time', 'endTime', 'DATETIME', '不可', '面接終了日時'],
            ['is_booked', 'isBooked', 'BOOLEAN', '不可', '予約済フラグ（default: false）'],
            ['meet_url', 'meetUrl', 'VARCHAR(500)', '可', 'Google Meet URL'],
            ['applicant_id', 'applicantId', 'INT (UNIQUE)', '可', '紐付け応募者ID（1:1リレーション）'],
            ['created_at', 'createdAt', 'DATETIME', '不可', '作成日時'],
            ['updated_at', 'updatedAt', 'DATETIME', '不可', '更新日時'],
        ]
    )
    doc.add_paragraph('')
    doc.add_paragraph('インデックス:')
    add_styled_table(doc,
        ['インデックス対象', '用途'],
        [
            ['start_time', '日時ベースのスロット検索'],
            ['is_booked', '空き枠フィルタリング'],
            ['applicant_id (UNIQUE)', '1:1リレーション制約'],
        ]
    )

    # 2.4 Applicant
    doc.add_heading('2.4 Applicant テーブル（応募者）', level=2)
    doc.add_paragraph('テーブル名: applicants')
    add_styled_table(doc,
        ['カラム名（DB）', 'Prismaフィールド名', 'データ型', 'NULL許可', '説明'],
        [
            ['id', 'id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['name', 'name', 'VARCHAR(100)', '不可', '氏名'],
            ['email', 'email', 'VARCHAR(255) UNIQUE', '不可', 'メールアドレス'],
            ['phone', 'phone', 'VARCHAR(20)', '可', '電話番号'],
            ['language', 'language', 'VARCHAR(5)', '不可', '言語設定（"ja" / "en"）'],
            ['job_category_id', 'jobCategoryId', 'INT (FK)', '不可', '応募職種ID → job_categories.id'],
            ['country_id', 'countryId', 'INT (FK)', '可', '国籍ID → countries.id'],
            ['visa_type_id', 'visaTypeId', 'INT (FK)', '可', '在留資格ID → visa_types.id'],
            ['postal_code', 'postalCode', 'VARCHAR(8)', '可', '郵便番号'],
            ['address', 'address', 'TEXT', '可', '住所'],
            ['building', 'building', 'VARCHAR(200)', '可', '建物名'],
            ['flow_status', 'flowStatus', 'ENUM', '不可', '応募フロー（default: INTERVIEW_WAITING）'],
            ['hiring_status', 'hiringStatus', 'ENUM', '不可', '採用ステータス（default: IN_PROGRESS）'],
            ['has_other_job', 'hasOtherJob', 'BOOLEAN', '可', '他の仕事をやっているか'],
            ['other_job_details', 'otherJobDetails', 'TEXT', '可', '他の仕事の詳細'],
            ['has_bank_in_japan', 'hasBankInJapan', 'BOOLEAN', '可', '日本の銀行口座の有無'],
            ['japanese_level', 'japaneseLevel', 'INT', '可', '日本語能力（1-5）'],
            ['english_level', 'englishLevel', 'INT', '可', '英語能力（1-5）'],
            ['communication_score', 'communicationScore', 'INT', '可', 'コミュニケーション評価（1-5）'],
            ['impression_score', 'impressionScore', 'INT', '可', '印象評価（1-5）'],
            ['interview_notes', 'interviewNotes', 'TEXT', '可', '面接備考欄'],
            ['created_at', 'createdAt', 'DATETIME', '不可', '作成日時（応募日時）'],
            ['updated_at', 'updatedAt', 'DATETIME', '不可', '更新日時'],
        ]
    )
    doc.add_paragraph('')
    doc.add_paragraph('インデックス:')
    add_styled_table(doc,
        ['インデックス対象', '用途'],
        [
            ['flow_status', 'フロー状態でのフィルタリング'],
            ['hiring_status', '採用ステータスでのフィルタリング'],
            ['created_at', '応募日時での範囲検索'],
        ]
    )

    # 2.5 リレーション
    doc.add_heading('2.5 リレーション図', level=2)
    doc.add_paragraph(
        'Applicant ←1:1→ InterviewSlot（applicantId UNIQUE FK）\n'
        'Applicant →N:1→ JobCategory（jobCategoryId FK）\n'
        'Applicant →N:1→ Country（countryId FK、既存テーブル）\n'
        'Applicant →N:1→ VisaType（visaTypeId FK、既存テーブル）'
    )

    # === 3. API設計 ===
    doc.add_heading('3. API設計', level=1)

    doc.add_heading('3.1 公開API（認証不要）', level=2)

    # 公開API一覧
    add_styled_table(doc,
        ['エンドポイント', 'メソッド', '説明'],
        [
            ['GET /api/interview-slots/available', 'GET', '未来の空きスロット一覧を返す'],
            ['POST /api/apply', 'POST', '応募情報送信 + スロット予約（トランザクション）+ 確認メール'],
            ['GET /api/job-categories/public', 'GET', 'アクティブな応募職種一覧'],
            ['GET /api/countries/public', 'GET', '国籍マスタ一覧'],
            ['GET /api/visa-types/public', 'GET', '在留資格マスタ一覧'],
        ]
    )

    doc.add_paragraph('')
    doc.add_heading('3.1.1 POST /api/apply（応募送信）の詳細', level=3 if hasattr(doc, '_') else 2)
    doc.add_paragraph('リクエストボディ:')
    add_styled_table(doc,
        ['フィールド', '型', '必須', '説明'],
        [
            ['name', 'string', '必須', '氏名'],
            ['email', 'string', '必須', 'メールアドレス'],
            ['phone', 'string', '任意', '電話番号'],
            ['language', 'string', '任意', '"ja"（デフォルト） / "en"'],
            ['jobCategoryId', 'number', '必須', '応募職種ID'],
            ['countryId', 'number', '任意', '国籍ID'],
            ['visaTypeId', 'number', '任意', '在留資格ID'],
            ['postalCode', 'string', '任意', '郵便番号'],
            ['address', 'string', '任意', '住所'],
            ['building', 'string', '任意', '建物名'],
            ['interviewSlotId', 'number', '必須', '選択した面接スロットID'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph('トランザクション処理フロー:')
    steps = [
        'メールアドレスの重複チェック（409 Conflict）',
        'スロットの空き確認 — isBooked=false かつ未来日時であること（409 Conflict）',
        'Applicant レコード作成',
        'InterviewSlot の applicantId 設定 + isBooked=true に更新',
        '監査ログ記録（actorType: SYSTEM）',
        'トランザクション外で確認メール非同期送信（失敗してもレスポンスは成功）',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {s}')

    doc.add_heading('3.2 管理者API（pms_session Cookie 必須）', level=2)
    add_styled_table(doc,
        ['エンドポイント', 'メソッド', '説明'],
        [
            ['GET /api/interview-slots', 'GET', '全スロット一覧（applicant情報含む）。?month=YYYY-MM で月別絞り込み'],
            ['POST /api/interview-slots', 'POST', 'スロット一括作成（日付 + 開始時間 + 終了時間 + 間隔指定）'],
            ['DELETE /api/interview-slots/[id]', 'DELETE', '空きスロット削除（予約済は拒否 400）'],
            ['GET /api/applicants', 'GET', '応募者一覧（ページネーション + flowStatus/hiringStatus/search フィルタ）'],
            ['GET /api/applicants/[id]', 'GET', '応募者詳細（全リレーション含む）'],
            ['PUT /api/applicants/[id]', 'PUT', '評価入力・ステータス変更（HIRED変更時に採用通知メール自動送信）'],
            ['GET /api/job-categories', 'GET', '全職種一覧（応募者数カウント含む）'],
            ['POST /api/job-categories', 'POST', '新規職種作成'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph('POST /api/interview-slots のリクエストボディ（一括生成モード）:')
    add_styled_table(doc,
        ['フィールド', '型', '必須', '説明'],
        [
            ['date', 'string', '必須', '日付（"2026-03-01"形式）'],
            ['startHour', 'number', '必須', '開始時間（9 = 9:00）'],
            ['endHour', 'number', '必須', '終了時間（17 = 17:00）'],
            ['intervalMinutes', 'number', '必須', 'スロット間隔（30 / 60）'],
            ['meetUrl', 'string', '任意', 'Google Meet URL（全スロット共通）'],
        ]
    )

    # === 4. 画面設計 ===
    doc.add_heading('4. 画面設計', level=1)

    doc.add_heading('4.1 応募フォーム（/apply）', level=2)
    add_styled_table(doc,
        ['項目', '内容'],
        [
            ['URL', '/apply'],
            ['アクセス権限', '認証不要（一般公開）'],
            ['サイドバー', '非表示（LayoutWrapper で除外）'],
            ['レスポンシブ', 'PC / スマートフォン両対応'],
            ['多言語対応', 'ja / en トグル切替（フォーム内の全ラベルが切替わる）'],
        ]
    )
    doc.add_paragraph('')
    doc.add_paragraph('フォーム構成:')
    sections = [
        '希望職種（JobCategory セレクトボックス）',
        '個人情報：氏名（必須）、メールアドレス + 確認用（一致チェック）、電話番号（自動フォーマット）',
        '国籍・在留資格：国籍ドロップダウン → 日本以外の場合にビザ種類表示（条件付き必須）',
        '住所：郵便番号（zipcloud API 自動補完）、住所、建物名',
        '面接日程選択：日付タブ + 時間帯カード形式。空きスロットのみ表示',
        '同意事項：プライバシーポリシー同意チェック + 入力内容確認チェック',
        '送信ボタン：全バリデーション通過後に有効化',
        '送信成功画面：面接日時・Google Meet URL を表示',
    ]
    for i, s in enumerate(sections, 1):
        doc.add_paragraph(f'{i}. {s}')

    doc.add_heading('4.2 管理者画面（/applicants）', level=2)
    add_styled_table(doc,
        ['項目', '内容'],
        [
            ['URL', '/applicants'],
            ['アクセス権限', '管理者（pms_session Cookie 保有者）のみ'],
            ['サイドバー', 'ORGANIZATION グループ内「応募者管理」（アイコン: bi-person-lines-fill）'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph('画面構成（タブ切替）:')

    doc.add_paragraph('タブ1: カレンダー', style='List Bullet')
    cal_items = [
        'FullCalendar ライブラリ使用（月表示 / 週表示トグル）',
        '空きスロット = 緑色、予約済みスロット = indigo 色で色分け表示',
        '日付クリック → スロット一括作成モーダルを表示（日付・時間帯・間隔・Meet URL 設定）',
        '予約済イベントクリック → 応募者評価モーダルを表示',
        '空きスロットクリック → 削除確認ダイアログ',
    ]
    for item in cal_items:
        doc.add_paragraph(f'  - {item}')

    doc.add_paragraph('タブ2: 応募者一覧', style='List Bullet')
    list_items = [
        '検索（氏名・メール）+ フロー/採用ステータスフィルタ',
        'テーブル列：氏名、メール、職種、面接日時、フローステータス、採用ステータス、操作',
        'ステータスは色付きバッジで表示',
        'ページネーション対応（20件/ページ）',
        '行クリック → 応募者評価モーダルを表示',
    ]
    for item in list_items:
        doc.add_paragraph(f'  - {item}')

    doc.add_heading('4.3 応募者評価モーダル', level=2)
    doc.add_paragraph('カレンダーまたは一覧から応募者をクリックすると表示されるモーダルダイアログ。')
    modal_sections = [
        '基本情報（読み取り専用）：氏名、メール、電話、言語、職種、面接日時、Google Meet リンクボタン',
        '国籍・住所（編集可能）：国籍、ビザ種類（外国籍の場合のみ）、郵便番号、住所、建物名',
        '面接評価（入力フォーム）：他の仕事（YES/NO + 詳細テキスト）、日本の口座有無、'
        '日本語能力（1-5）、英語能力（1-5）※外国籍のみ、コミュニケーション（1-5）、印象（1-5）、備考欄',
        'ステータス変更：フローステータス、採用ステータス。HIRED 選択時に採用通知メール送信の警告表示',
    ]
    for i, s in enumerate(modal_sections, 1):
        doc.add_paragraph(f'{i}. {s}')

    # === 5. メール設計 ===
    doc.add_heading('5. メール設計', level=1)
    doc.add_paragraph('実装ファイル: src/lib/mailer.ts')

    add_styled_table(doc,
        ['関数名', 'トリガー', '言語', '内容'],
        [
            ['sendApplicantConfirmationEmail', 'POST /api/apply 成功後', 'ja/en', '面接予約確認。職種名・面接日時・Google Meet URL を記載'],
            ['sendHiringNotificationEmail', 'PUT /api/applicants/[id] で hiringStatus→HIRED', 'ja/en', '採用通知。職種名・お祝いメッセージ・今後の手続き案内'],
        ]
    )
    doc.add_paragraph('')
    doc.add_paragraph('送信パターン: 非同期（.catch でエラーキャッチ）。メール送信失敗時もAPI処理は成功扱い。')

    # === 6. 監査ログ連携 ===
    doc.add_heading('6. 監査ログ連携', level=1)
    add_styled_table(doc,
        ['対象API', 'アクター種別', 'アクション', 'tx渡し', '備考'],
        [
            ['POST /api/apply', 'SYSTEM', 'CREATE', 'あり', '応募者が自ら応募（認証なし）'],
            ['POST /api/interview-slots', 'EMPLOYEE', 'CREATE', 'あり', '管理者によるスロット一括作成'],
            ['DELETE /api/interview-slots/[id]', 'EMPLOYEE', 'DELETE', 'あり', '管理者によるスロット削除'],
            ['PUT /api/applicants/[id]', 'EMPLOYEE', 'UPDATE / STATUS_CHANGE', 'あり', '評価入力またはステータス変更'],
            ['POST /api/job-categories', 'EMPLOYEE', 'CREATE', 'あり', '管理者による職種作成'],
        ]
    )

    # === 7. ファイル一覧 ===
    doc.add_heading('7. ファイル一覧', level=1)
    add_styled_table(doc,
        ['ファイルパス', '種別', '説明'],
        [
            ['prisma/schema.prisma', '変更', 'Enum 2種 + Model 3種 + 逆リレーション2箇所追加'],
            ['src/middleware.ts', '変更', '公開パス（/apply, /api/apply 等）を isPublicPath に追加'],
            ['src/components/LayoutWrapper.tsx', '変更', '/apply でサイドバー非表示'],
            ['src/components/Sidebar.tsx', '変更', 'ORGANIZATION グループに「応募者管理」追加'],
            ['src/lib/mailer.ts', '変更', 'メール関数 2種追加'],
            ['src/app/apply/page.tsx', '新規', '応募者向け公開フォーム'],
            ['src/app/applicants/page.tsx', '新規', '管理者向けカレンダー＋応募者一覧＋評価モーダル'],
            ['src/app/api/apply/route.ts', '新規', '応募送信API（公開）'],
            ['src/app/api/interview-slots/available/route.ts', '新規', '空きスロットAPI（公開）'],
            ['src/app/api/interview-slots/route.ts', '新規', 'スロット一覧・作成API（管理者）'],
            ['src/app/api/interview-slots/[id]/route.ts', '新規', 'スロット削除API（管理者）'],
            ['src/app/api/applicants/route.ts', '新規', '応募者一覧API（管理者）'],
            ['src/app/api/applicants/[id]/route.ts', '新規', '応募者詳細・更新API（管理者）'],
            ['src/app/api/job-categories/route.ts', '新規', '職種マスタ管理API（管理者）'],
            ['src/app/api/job-categories/public/route.ts', '新規', '職種一覧API（公開）'],
            ['src/app/api/countries/public/route.ts', '新規', '国籍一覧API（公開）'],
            ['src/app/api/visa-types/public/route.ts', '新規', 'ビザ種類一覧API（公開）'],
        ]
    )

    # === 8. 依存パッケージ ===
    doc.add_heading('8. 追加依存パッケージ', level=1)
    add_styled_table(doc,
        ['パッケージ名', 'バージョン', '用途'],
        [
            ['@fullcalendar/react', '^6.x', 'カレンダーUI（React ラッパー）'],
            ['@fullcalendar/core', '^6.x', 'FullCalendar コアエンジン'],
            ['@fullcalendar/daygrid', '^6.x', '月表示プラグイン'],
            ['@fullcalendar/timegrid', '^6.x', '週表示（時間軸）プラグイン'],
            ['@fullcalendar/interaction', '^6.x', 'ドラッグ・クリック操作プラグイン'],
        ]
    )

    # === 9. 今後の拡張候補 ===
    doc.add_heading('9. 今後の拡張候補', level=1)
    future = [
        'Google Calendar API 連携による Meet URL 自動生成',
        '面接リマインダーメールの自動送信（面接前日）',
        '応募者への不採用通知メール',
        '応募統計ダッシュボード（職種別・月別応募数グラフ）',
        '応募フォームのカスタマイズ（管理者が項目を追加・非表示にする機能）',
    ]
    for f in future:
        doc.add_paragraph(f, style='List Bullet')

    # 保存
    path = os.path.join(OUTPUT_DIR, 'PMS_応募者管理機能_設計書.docx')
    doc.save(path)
    print(f"[OK] 設計書を保存: {path}")
    return path


# ====================================================================
# 2. 機能一覧に応募者管理セクション追加
# ====================================================================

def update_feature_list():
    path = os.path.join(OUTPUT_DIR, 'PMS_機能一覧.docx')
    doc = Document(path)

    # 更新日を変更
    for p in doc.paragraphs:
        if '最終更新' in p.text:
            for run in p.runs:
                if '最終更新' in run.text:
                    run.text = run.text.replace('2026年02月26日', '2026年02月27日').replace('1.0', '1.1')
                    break

    # 最後のHeading 1の番号を取得して次の番号を計算
    last_heading_num = 0
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1' and p.text and p.text[0].isdigit():
            try:
                num = int(p.text.split('.')[0])
                if num > last_heading_num:
                    last_heading_num = num
            except ValueError:
                pass

    next_num = last_heading_num + 1

    # 新セクション追加
    doc.add_heading(f'{next_num}. 応募者管理・面接日程調整機能  /  Applicant Management & Interview Scheduling', level=1)
    doc.add_paragraph('応募者情報の管理、面接スロットの作成・予約、面接評価・採用プロセスを一元管理します。')

    # 応募者管理テーブル
    headers = ['カテゴリ / Category', '機能名（日本語）', '機能名（English）', 'URL / エンドポイント', '概要', '稼働ステータス\nStatus']

    rows_data = [
        ['応募フォーム\nApplication Form', '応募フォーム（日/英）', 'Job Application Form (JA/EN)', '/apply  /  POST /api/apply',
         '応募者が職種選択・個人情報入力・面接枠選択・同意チェックを行い応募を送信。ja/en言語切替対応', ''],
        ['', '面接枠表示・予約', 'Interview Slot Selection', 'GET /api/interview-slots/available',
         '未来の空きスロットを日付別に表示し、応募者が1つ選択。トランザクションで競合防止', ''],
        ['', '応募確認メール送信', 'Application Confirmation Email', 'POST /api/apply（内部処理）',
         '応募完了後に面接日時・Meet URLを記載した確認メールを自動送信（ja/en対応）', ''],
        ['', '郵便番号自動補完', 'Postal Code Auto-Complete', 'zipcloud API',
         '郵便番号7桁入力時にzipcloud APIから住所を自動取得・入力', ''],
        ['面接スロット管理\nInterview Slot Management', '面接カレンダー表示', 'Interview Calendar View', '/applicants（カレンダータブ）',
         'FullCalendarで月/週表示。空き枠=緑、予約済=indigoで色分け', ''],
        ['', 'スロット一括作成', 'Bulk Slot Creation', 'POST /api/interview-slots',
         '日付・開始/終了時間・間隔（30/60分）・Meet URLを指定して一括作成', ''],
        ['', 'スロット削除', 'Slot Deletion', 'DELETE /api/interview-slots/[id]',
         '未予約のスロットを削除（予約済みは削除不可）', ''],
        ['応募者管理\nApplicant Management', '応募者一覧・検索', 'Applicant List & Search', '/applicants（一覧タブ）  /  GET /api/applicants',
         '氏名・メール検索、フロー/採用ステータスフィルタ、ページネーション対応', ''],
        ['', '応募者詳細・評価入力', 'Applicant Detail & Evaluation', 'GET/PUT /api/applicants/[id]',
         '基本情報表示、国籍・住所修正、語学力・コミュニケーション・印象（1-5）評価入力', ''],
        ['', 'フローステータス変更', 'Flow Status Change', 'PUT /api/applicants/[id]',
         '面接待ち→研修待ち→研修完了のフロー管理', ''],
        ['', '採用ステータス変更', 'Hiring Status Change', 'PUT /api/applicants/[id]',
         '選考中→採用/不採用の判定。採用時に通知メール自動送信', ''],
        ['', '採用通知メール送信', 'Hiring Notification Email', 'PUT /api/applicants/[id]（内部処理）',
         'hiringStatus=HIREDへの変更時に採用お祝いメールを自動送信（ja/en対応）', ''],
        ['職種マスタ管理\nJob Category Master', '職種一覧・作成', 'Job Category CRUD', '/applicants（職種マスタモーダル）  /  /api/job-categories',
         '応募職種のマスタデータ管理。日本語/英語名、有効/無効フラグ、表示順設定', ''],
        ['', '公開職種API', 'Public Job Categories API', 'GET /api/job-categories/public',
         '応募フォーム向けにアクティブな職種一覧を返す公開API', ''],
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行（既存テーブルと同じスタイル）
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    # データ行
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    doc.save(path)
    print(f"[OK] 機能一覧を更新: {path}")
    return path


# ====================================================================
# 実行
# ====================================================================
if __name__ == '__main__':
    create_design_doc()
    update_feature_list()
    print("\n完了！")
