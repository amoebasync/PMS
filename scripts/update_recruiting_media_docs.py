#!/usr/bin/env python3
"""
求人媒体トラッキング・国エイリアス・面接管理改修 設計書更新スクリプト
- PMS_応募者管理機能_設計書.docx に新機能セクションを追記
- PMS_機能一覧.docx に項目を追加
"""
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = os.path.expanduser("~/Downloads/PMS_設計書")


def set_cell_shading(cell, color_hex):
    """セルの背景色を設定"""
    shading_elm = docx.oxml.parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, rows):
    """ヘッダー付きテーブルを追加"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1E293B")

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F8FAFC")

    return table


# ====================================================================
# 1. 設計書の更新
# ====================================================================

def update_design_doc():
    path = os.path.join(OUTPUT_DIR, 'PMS_応募者管理機能_設計書.docx')
    doc = Document(path)

    # バージョンを1.1に更新
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '1.0' in cell.text and '文書バージョン' not in cell.text:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if '1.0' in run.text:
                                run.text = run.text.replace('1.0', '1.1')

    # === 新セクション追加 ===

    # 10. 求人媒体トラッキング機能
    doc.add_page_break()
    doc.add_heading('10. 求人媒体トラッキング機能（2026-02-28 追加）', level=1)

    doc.add_heading('10.1 概要', level=2)
    doc.add_paragraph(
        '応募フォーム（/apply）のURLパラメータ ?source=xxx で求人媒体を自動追跡する機能。'
        '管理者は設定画面で求人媒体マスタを管理し、応募者管理画面の評価モーダルで応募経路を確認・変更できる。'
    )

    doc.add_heading('10.2 RecruitingMedia テーブル', level=2)
    doc.add_paragraph('テーブル名: recruiting_media')
    add_styled_table(doc,
        ['カラム名（DB）', 'Prismaフィールド名', 'データ型', 'NULL許可', '説明'],
        [
            ['id', 'id', 'INT AUTO_INCREMENT', '不可', '主キー'],
            ['name_ja', 'nameJa', 'VARCHAR(100)', '不可', '媒体名（日本語）'],
            ['name_en', 'nameEn', 'VARCHAR(100)', '可', '媒体名（英語）'],
            ['code', 'code', 'VARCHAR(50) UNIQUE', '不可', 'URLパラメータ値（例: indeed, mynavi）'],
            ['is_active', 'isActive', 'BOOLEAN', '不可', '有効フラグ（default: true）'],
            ['sort_order', 'sortOrder', 'INT', '不可', '表示順（default: 100）'],
            ['created_at', 'createdAt', 'DATETIME', '不可', '作成日時'],
            ['updated_at', 'updatedAt', 'DATETIME', '不可', '更新日時'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph('Applicant テーブルへの追加カラム:')
    add_styled_table(doc,
        ['カラム名（DB）', 'Prismaフィールド名', 'データ型', 'NULL許可', '説明'],
        [
            ['recruiting_media_id', 'recruitingMediaId', 'INT (FK)', '可', '求人媒体ID → recruiting_media.id'],
        ]
    )

    doc.add_heading('10.3 API', level=2)
    add_styled_table(doc,
        ['エンドポイント', 'メソッド', '認証', '説明'],
        [
            ['/api/recruiting-media', 'GET', 'pms_session', '媒体一覧（応募者数カウント含む）'],
            ['/api/recruiting-media', 'POST', 'pms_session', '媒体作成（code自動小文字化、重複チェック）'],
            ['/api/recruiting-media?id=X', 'PUT', 'pms_session', '媒体更新'],
            ['/api/recruiting-media?id=X', 'DELETE', 'pms_session', '媒体削除（応募者紐付き時は拒否）'],
            ['/api/recruiting-media/public?code=xxx', 'GET', '不要', 'codeからIDを解決（応募フォーム用）'],
        ]
    )

    doc.add_heading('10.4 トラッキングフロー', level=2)
    steps = [
        '管理者が設定画面の「求人媒体」タブで媒体を登録（例: code=indeed, nameJa=Indeed）',
        '求人広告に /apply?source=indeed のURLを掲載',
        '応募者がそのURLからアクセスすると、フロントエンドがGET /api/recruiting-media/public?code=indeedを呼び出し',
        '取得した媒体IDをstateに保存（UIには非表示）',
        '応募送信時にPOST /api/applyのbodyにrecruitingMediaIdを含めて送信',
        'DBのapplicants.recruiting_media_idに保存',
        '管理者は /applicants の評価モーダルで「応募経路」ドロップダウンから確認・変更可能',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {s}')

    # 11. 国名エイリアス機能
    doc.add_heading('11. 国名エイリアス機能（2026-02-28 追加）', level=1)

    doc.add_heading('11.1 概要', level=2)
    doc.add_paragraph(
        '国マスタにaliases（別名）フィールドを追加し、応募フォームの国検索で日本語・英語・通称すべてでマッチするようにする。'
    )

    doc.add_heading('11.2 データモデル変更', level=2)
    doc.add_paragraph('Country テーブルへの追加カラム:')
    add_styled_table(doc,
        ['カラム名（DB）', 'Prismaフィールド名', 'データ型', 'NULL許可', '説明'],
        [
            ['aliases', 'aliases', 'TEXT', '可', 'カンマ区切りの別名（例: 韓国,大韓民国,Republic of Korea,South Korea）'],
        ]
    )

    doc.add_heading('11.3 検索ロジック', level=2)
    doc.add_paragraph(
        '応募フォームの国選択（SearchableSelect）でfilterFn propを使用し、以下の順にマッチ:\n'
        '1. country.name（日本語名）\n'
        '2. country.nameEn（英語名）\n'
        '3. country.aliases（別名、カンマ区切りを分割して各要素と照合）\n'
        '\n'
        'SearchableSelectコンポーネントにfilterFn?: (option: T, search: string) => boolean プロパティを追加し、'
        'カスタムフィルタロジックを外部から注入可能にした。'
    )

    doc.add_heading('11.4 初期データ投入', level=2)
    doc.add_paragraph(
        'スクリプト: scripts/seed-country-aliases.ts\n'
        '実行: npx tsx scripts/seed-country-aliases.ts\n'
        '対象: 34カ国の別名を一括登録（既にエイリアス設定済みの国はスキップ）'
    )

    # 12. 管理者面接キャンセル・日程変更機能
    doc.add_heading('12. 管理者面接キャンセル・日程変更機能（2026-02-28 追加）', level=1)

    doc.add_heading('12.1 概要', level=2)
    doc.add_paragraph(
        '管理者が応募者管理画面の評価モーダルから面接のキャンセルおよび日程変更を行える機能。'
    )

    doc.add_heading('12.2 API', level=2)
    add_styled_table(doc,
        ['エンドポイント', 'メソッド', '認証', '説明'],
        [
            ['POST /api/applicants/[id]/cancel-interview', 'POST', 'pms_session',
             '面接キャンセル。スロット解放（isBooked=false, applicantId=null）。応募者ステータスは変更しない'],
            ['POST /api/applicants/[id]/reschedule', 'POST', 'pms_session',
             '面接日程変更。body: { newSlotId }。旧スロット解放→新スロット予約→Google Meet自動生成'],
        ]
    )

    doc.add_heading('12.3 面接キャンセル処理フロー', level=2)
    cancel_steps = [
        '応募者+interviewSlotを取得',
        'スロットが紐付いていない場合はエラー（400）',
        'トランザクション: スロット解放（isBooked=false, applicantId=null）',
        '監査ログ記録（actorType: EMPLOYEE, action: UPDATE）',
        '応募者のステータスは変更しない（管理者が手動で判断）',
    ]
    for i, s in enumerate(cancel_steps, 1):
        doc.add_paragraph(f'{i}. {s}')

    doc.add_heading('12.4 面接日程変更処理フロー', level=2)
    reschedule_steps = [
        '応募者+現在のinterviewSlotを取得',
        '新スロットの空き確認・未来日チェック',
        'トランザクション:',
        '  a. 旧スロット解放（isBooked=false, applicantId=null）',
        '  b. 新スロット予約（isBooked=true, applicantId設定）',
        '  c. Google Meet自動生成（設定済みの場合のみ）',
        '  d. 監査ログ記録',
        '更新後の応募者+スロット情報を返却',
    ]
    for i, s in enumerate(reschedule_steps, 1):
        doc.add_paragraph(f'{i}. {s}')

    doc.add_heading('12.5 UI', level=2)
    doc.add_paragraph(
        '評価モーダルの面接情報表示エリアに「日程変更」「面接キャンセル」ボタンを追加。\n'
        '- 日程変更: パネルが開き、空きスロット一覧から新しい日程を選択\n'
        '- キャンセル: 確認ダイアログ後に実行\n'
        '- 評価スケール（ScoreSelector）: 低/高ラベル付き、1-2=赤系、3=黄系、4-5=緑系のカラーコーディング\n'
        '- Google Meetボタン: テキストリンクスタイルに変更（背景色なし、アイコン小型化）'
    )

    # 13. 追加ファイル一覧
    doc.add_heading('13. 追加ファイル一覧（2026-02-28）', level=1)
    add_styled_table(doc,
        ['ファイルパス', '種別', '説明'],
        [
            ['prisma/schema.prisma', '変更', 'RecruitingMedia モデル追加、Applicant.recruitingMediaId追加、Country.aliases追加'],
            ['src/middleware.ts', '変更', '/api/recruiting-media/public を公開パスに追加'],
            ['src/app/api/recruiting-media/route.ts', '新規', '求人媒体 CRUD API（管理者）'],
            ['src/app/api/recruiting-media/public/route.ts', '新規', '求人媒体 公開API（code→ID解決）'],
            ['src/app/api/applicants/[id]/cancel-interview/route.ts', '新規', '面接キャンセルAPI'],
            ['src/app/api/applicants/[id]/reschedule/route.ts', '新規', '面接日程変更API'],
            ['src/app/api/apply/route.ts', '変更', 'recruitingMediaId 受け取り追加'],
            ['src/app/api/applicants/route.ts', '変更', 'recruitingMedia include追加'],
            ['src/app/api/applicants/[id]/route.ts', '変更', 'recruitingMedia include + 更新対応'],
            ['src/app/api/countries/public/route.ts', '変更', 'aliases フィールド追加'],
            ['src/app/settings/page.tsx', '変更', '求人媒体タブ追加 + 国エイリアス入力欄'],
            ['src/app/apply/page.tsx', '変更', '媒体トラッキング + filterFn + 国検索改善'],
            ['src/app/applicants/page.tsx', '変更', '媒体ドロップダウン + キャンセル/変更 + ScoreSelector + Meetボタン'],
            ['scripts/seed-country-aliases.ts', '新規', '国エイリアス初期データ投入スクリプト'],
        ]
    )

    doc.save(path)
    print(f"[OK] 設計書を更新: {path}")
    return path


# ====================================================================
# 2. 機能一覧の更新
# ====================================================================

def update_feature_list():
    path = os.path.join(OUTPUT_DIR, 'PMS_機能一覧.docx')
    doc = Document(path)

    # 更新日を変更
    for p in doc.paragraphs:
        if '最終更新' in p.text:
            for run in p.runs:
                if '最終更新' in run.text:
                    run.text = run.text.replace('2026年02月27日', '2026年02月28日')
                    break

    # 最後のHeading 1の番号を取得して次の番号を計算
    last_heading_num = 0
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1' and p.text and p.text[0].isdigit():
            try:
                num = int(p.text.split('.')[0])
                if num > last_heading_num:
                    last_heading_num = num
            except ValueError:
                pass

    next_num = last_heading_num + 1

    # 新セクション追加
    doc.add_heading(f'{next_num}. 求人媒体トラッキング・国名エイリアス・面接管理改修  /  Recruiting Media Tracking & Country Aliases & Interview Management', level=1)
    doc.add_paragraph('求人媒体のURLパラメータトラッキング、国検索の多言語化、管理画面での面接キャンセル/日程変更、評価UIの改善を含む。')

    headers = ['カテゴリ / Category', '機能名（日本語）', '機能名（English）', 'URL / エンドポイント', '概要', '稼働ステータス\nStatus']

    rows_data = [
        ['求人媒体トラッキング\nRecruiting Media Tracking', '求人媒体マスタ管理', 'Recruiting Media Master CRUD', '/settings（求人媒体タブ）\n/api/recruiting-media',
         '求人媒体の登録・編集・削除。code（URLパラメータ値）、名称（日英）、表示順、有効/無効を管理', ''],
        ['', '応募経路自動追跡', 'Source Auto-Tracking via URL', '/apply?source=xxx\nGET /api/recruiting-media/public?code=xxx',
         '応募フォームのURLパラメータで求人媒体を自動追跡。UIには非表示', ''],
        ['', '応募経路確認・変更', 'Source Confirmation & Change', '/applicants（評価モーダル）',
         '応募者管理画面の評価モーダルで応募経路ドロップダウンから確認・変更可能', ''],
        ['国名エイリアス\nCountry Aliases', '国名別名管理', 'Country Alias Management', '/settings（国籍タブ）',
         '国マスタの作成/編集画面にaliases（別名）入力欄を追加。カンマ区切りで複数の別名を登録', ''],
        ['', '多言語国検索', 'Multilingual Country Search', '/apply（国選択）',
         '応募フォームの国検索で日本語名・英語名・別名すべてにマッチ。SearchableSelectにfilterFn propを追加', ''],
        ['', '国エイリアス初期データ投入', 'Country Alias Seed', 'npx tsx scripts/seed-country-aliases.ts',
         '主要34カ国の別名を一括投入するスクリプト', ''],
        ['面接管理改修\nInterview Management Enhancement', '面接キャンセル（管理者）', 'Admin Interview Cancel', 'POST /api/applicants/[id]/cancel-interview',
         '管理者が評価モーダルから面接をキャンセル。スロット解放のみでステータスは変更しない', ''],
        ['', '面接日程変更（管理者）', 'Admin Interview Reschedule', 'POST /api/applicants/[id]/reschedule',
         '管理者が評価モーダルから面接日程を変更。旧スロット解放→新スロット予約→Google Meet自動生成', ''],
        ['', '評価スケール改善', 'Score Selector Enhancement', '/applicants（評価モーダル）',
         'ScoreSelectorに低/高ラベル追加。1-2=赤系、3=黄系、4-5=緑系のカラーコーディング', ''],
        ['', 'Google Meet ボタン改善', 'Meet Button Style Update', '/applicants（評価モーダル）',
         'Google Meetボタンをテキストリンクスタイルに変更（背景色なし、コンパクト化）', ''],
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    doc.save(path)
    print(f"[OK] 機能一覧を更新: {path}")
    return path


# ====================================================================
# 実行
# ====================================================================
if __name__ == '__main__':
    update_design_doc()
    update_feature_list()
    print("\n完了！")
